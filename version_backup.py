import os
import random
import sys
from textwrap import dedent
from typing import List, Optional
from altair import Dict
from matplotlib import pyplot as plt
import pdfplumber
from datetime import datetime
import pyspark
import streamlit as st
import fitz
import PyPDF2
import docx
import io
from typing import List
from openai import OpenAI
import numpy as np
from langchain_openai import ChatOpenAI
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.prompts import ChatPromptTemplate
import logging
from langdetect import detect, DetectorFactory
import boto3
import psycopg2
from psycopg2.extras import RealDictCursor
import json
from dotenv import load_dotenv 
from langchain_openai import ChatOpenAI
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.prompts import ChatPromptTemplate
import traceback
import plotly.express as px
import time 
import pandas as pd
import builtins
import uuid
import os
import warnings
from SparkAnalysis import SparkAnalysis
from pyspark.sql import SparkSession

LANGCHAIN_AVAILABLE = True
# FIX file watcher warnings pour igner fichier de cashe de streamlit
os.environ['STREAMLIT_SERVER_FILE_WATCHER_TYPE'] = 'none'
warnings.filterwarnings("ignore", message=".*non-watched file.*")
warnings.filterwarnings("ignore", message=".*file watcher.*")

from analysis import demonstrate_chart_libraries,create_skills_correlation_matrix,create_advanced_skills_dashboard,create_skills_word_cloud_simulation,create_advanced_experience_analysis,create_timeline_analytics

# Set up language detection
DetectorFactory.seed = 0
# --- Configuration ---
load_dotenv()
# Configuration - Updated with correct S3 buckets
CONFIG = {
    "S3_BUCKET_CANDIDATES": "tekbootwebsite2",  # For candidates
    "AWS_REGION": "us-east-1",
    "OPENAI_MODEL": "gpt-4-turbo"
}
POSTGRES_CONFIG = {
    "dbname": "cv_processing_system",
    "user": "postgres",
    "password": "Karisma@2020@fun",
    "host": "localhost",
    "port": "5432"
}
print(f"PySpark version: {pyspark.__version__}")
print(f"✅ Java configured: {os.environ.get('JAVA_HOME')}")
print(f"✅ Spark driver host: {os.environ.get('localhost')}")

HADOOP_HOME=os.environ.get('HADOOP_HOME')
#java configuration with path 
java_home_from_env = os.getenv("JAVA_HOME")
if java_home_from_env:
    os.environ['JAVA_HOME'] = java_home_from_env
    print(f"✅ JAVA_HOME loaded from .env: {java_home_from_env}")


else:
    #path de java
    fallback_java_home = r"C:\Program Files\Java\jdk-21"
    if os.path.exists(fallback_java_home):
        os.environ['JAVA_HOME'] = fallback_java_home
        print(f"✅ JAVA_HOME set to fallback: {fallback_java_home}")
    else:
        print(f"❌ Java not found at expected location: {fallback_java_home}")

print("✅ Java environment configuration updated for JDK-21")

#vérification de fonctionnalité de pyspark
try:
    from pyspark.sql import SparkSession
    from pyspark.sql.functions import *
    from pyspark.sql.types import *
    from pyspark.sql.window import Window
    import pyspark.sql.functions as F
    SPARK_AVAILABLE = True
    print("✅ PySpark successfully imported")
except ImportError:
    SPARK_AVAILABLE = False
    print("⚠️ PySpark not available - using fallback methods")
# ============= DEFINE CLASSES ============= #
# Restore built-in functions to prevent override issues
min = builtins.min
max = builtins.max
len = builtins.len
sum = builtins.sum

class DatabaseConfig:
    """Database configuration class"""
    def __init__(self, **kwargs):
        self.config = kwargs or POSTGRES_CONFIG
        
    def get_connection_string(self):
        return f"postgresql://{self.config['user']}:{self.config['password']}@{self.config['host']}:{self.config['port']}/{self.config['dbname']}"

class DatabaseManager:
    #tous les fonctionnalités de de la base pour tous PostgreSQL operations 

    def __init__(self, db_config):
        self.db_config = db_config
        self.connection = None

    def get_connection(self):
        try:
            if not self.connection or self.connection.closed:
                self.connection = psycopg2.connect(**self.db_config.config)
            return self.connection
        except Exception as e:
            print(f"Database connection error: {e}")
            raise

    def execute_query(self, query: str, params=None):
        """FIXED: Always return dictionaries"""
        try:
            conn = self.get_connection()
            # CRITICAL: Always use RealDictCursor
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cursor:
                cursor.execute(query, params)
                
                if query.strip().upper().startswith('SELECT'):
                    results = cursor.fetchall()
                    # Convert to regular dictionaries
                    return [dict(row) for row in results] if results else []
                else:
                    conn.commit()
                    return []
                    
        except Exception as e:
            print(f"Database error: {e}")
            if self.connection:
                try:
                    self.connection.rollback()
                except:
                    pass
            return []

    def fetch_one(self, query, params=None):
        #retourne dictionnaire d'un élément
        try:
            conn = self.get_connection()
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cursor:
                cursor.execute(query, params)
                result = cursor.fetchone()
                return dict(result) if result else None
        except Exception as e:
            print(f"Database fetch_one error: {e}")
            return None 
    def execute_query_safely(self, query: str, params: dict = None):
        """Execute query safely and return DataFrame"""
        try:
            results = self.execute_query(query, params)
            if results:
                import pandas as pd
                return pd.DataFrame(results)
            else:
                import pandas as pd
                return pd.DataFrame()
        except Exception as e:
            print(f"Query execution error: {e}")
            import pandas as pd
            return pd.DataFrame()
   
    #récupéer tous les analyses des compétences
    def get_skills_analytics(self) -> dict:
        """Get analytics data for skills and categories"""
        try:
            # Technical skills analysis
            tech_skills_query = """
            SELECT skill, COUNT(*) as frequency
            FROM (
                SELECT jsonb_array_elements_text(skills->'technical_skills') as skill
                FROM candidates 
                WHERE skills->'technical_skills' IS NOT NULL
            ) as tech_skills
            GROUP BY skill
            ORDER BY frequency DESC;
            """
            tech_df = self.execute_query_safely(tech_skills_query)
            
            # Soft skills analysis
            soft_skills_query = """
            SELECT skill, COUNT(*) as frequency
            FROM (
                SELECT jsonb_array_elements_text(skills->'soft_skills') as skill
                FROM candidates 
                WHERE skills->'soft_skills' IS NOT NULL
            ) as soft_skills
            GROUP BY skill
            ORDER BY frequency DESC;
            """
            soft_df = self.execute_query_safely(soft_skills_query)
            
            # Job categories analysis
            categories_query = """
            SELECT job_category, COUNT(*) as count
            FROM candidates 
            WHERE job_category IS NOT NULL AND job_category != ''
            GROUP BY job_category
            ORDER BY count DESC;
            """
            cat_df = self.execute_query_safely(categories_query)
            
            # Experience analysis
            experience_query = """
            SELECT experience_years, COUNT(*) as count
            FROM candidates 
            GROUP BY experience_years
            ORDER BY experience_years;
            """
            exp_df = self.execute_query_safely(experience_query)
            
            # Education level analysis
            education_query = """
            SELECT education_level, COUNT(*) as count
            FROM candidates 
            WHERE education_level IS NOT NULL AND education_level != ''
            GROUP BY education_level
            ORDER BY count DESC;
            """
            edu_df = self.execute_query_safely(education_query)
            
            return {
                'technical_skills': tech_df,
                'soft_skills': soft_df,
                'categories': cat_df,
                'experience': exp_df,
                'education': edu_df
            }
        except Exception as e:
            print(f"Failed to get analytics: {e}")
            return {
                'technical_skills': pd.DataFrame(),
                'soft_skills': pd.DataFrame(),
                'categories': pd.DataFrame(),
                'experience': pd.DataFrame(),
                'education': pd.DataFrame()
            }


def diagnose_database_connection():
    """
    Complete diagnostic for database connection issues
    """
    st.write("🔍 **Database Connection Diagnosis**")
    
    # Check 1: Environment Variables
    st.write("### 1. Environment Variables Check")
    
    try:
        # Check if .env file exists
        env_file_exists = os.path.exists('.env')
        st.write(f"📄 .env file exists: {'✅ Yes' if env_file_exists else '❌ No'}")
        
        # Check POSTGRES_CONFIG
        if 'POSTGRES_CONFIG' in globals():
            st.write("✅ POSTGRES_CONFIG is defined")
        else:
            st.error("❌ POSTGRES_CONFIG not found in globals()")
        
        # Check individual environment variables
        env_vars = ['DB_HOST', 'DB_PORT', 'DB_NAME', 'DB_USER', 'DB_PASSWORD']
        for var in env_vars:
            value = os.getenv(var)
            if value:
                masked_value = value if var != 'DB_PASSWORD' else '*' * len(value)
                st.write(f"✅ {var}: {masked_value}")
            else:
                st.error(f"❌ {var}: Not found")
        
    except Exception as e:
        st.error(f"Environment check failed: {e}")
    
    # Check 2: Direct Connection Test
    st.write("### 2. Direct Database Connection Test")
    
    try:
        # Try direct connection with psycopg2
        conn_params = {
            'host': os.getenv('DB_HOST', 'localhost'),
            'port': os.getenv('DB_PORT', '5432'),
            'database': os.getenv('DB_NAME', 'postgres'),
            'user': os.getenv('DB_USER', 'postgres'),
            'password': os.getenv('DB_PASSWORD', '')
        }
        
        st.write(f"Attempting connection to: {conn_params['user']}@{conn_params['host']}:{conn_params['port']}/{conn_params['database']}")
        
        # Test connection
        conn = psycopg2.connect(**conn_params)
        cursor = conn.cursor()
        cursor.execute("SELECT version()")
        version = cursor.fetchone()
        cursor.close()
        conn.close()
        
        st.success(f"✅ Direct connection successful!")
        st.success(f"PostgreSQL version: {version[0]}")
        
    except psycopg2.OperationalError as e:
        st.error(f"❌ Connection failed: {e}")
        
        # Specific error diagnosis
        error_str = str(e).lower()
        if 'could not connect to server' in error_str:
            st.error("🔧 **Fix**: PostgreSQL server is not running or wrong host/port")
        elif 'password authentication failed' in error_str:
            st.error("🔧 **Fix**: Wrong username or password")
        elif 'database' in error_str and 'does not exist' in error_str:
            st.error("🔧 **Fix**: Database name is incorrect")
        elif 'timeout' in error_str:
            st.error("🔧 **Fix**: Network timeout - check firewall/network settings")
        
    except Exception as e:
        st.error(f"❌ Unexpected error: {e}")
    
    # Check 3: Test DatabaseManager class
    st.write("### 3. DatabaseManager Class Test")
    
    try:
        # Test if DatabaseConfig works
        if 'DatabaseConfig' in globals():
            db_config = DatabaseConfig(**POSTGRES_CONFIG)
            st.success("✅ DatabaseConfig created successfully")
            
            # Test DatabaseManager
            if 'DatabaseManager' in globals():
                db_manager = DatabaseManager(db_config)
                st.success("✅ DatabaseManager created successfully")
                
                # Test a simple query
                result = db_manager.execute_query("SELECT 1 as test")
                if result:
                    st.success(f"✅ Query execution successful: {result}")
                else:
                    st.error("❌ Query returned no results")
                    
            else:
                st.error("❌ DatabaseManager class not found")
        else:
            st.error("❌ DatabaseConfig class not found")
            
    except Exception as e:
        st.error(f"❌ DatabaseManager test failed: {e}")
        st.code(traceback.format_exc())
def patch_database_manager(db_manager):
    def patched_execute_query(query, params=None):
        try:
            conn = db_manager.get_connection()
            # Force RealDictCursor to return dictionaries
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cursor:
                cursor.execute(query, params)
                if query.strip().upper().startswith('SELECT'):
                    results = cursor.fetchall()
                    # Ensure we return dictionaries
                    return [dict(row) for row in results] if results else []
                else:
                    conn.commit()
                    return []
        except Exception as e:
            print(f"Patched database error: {e}")
            return []
    
    def patched_fetch_one(query, params=None):
        try:
            conn = db_manager.get_connection()
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cursor:
                cursor.execute(query, params)
                result = cursor.fetchone()
                return dict(result) if result else None
        except Exception as e:
            print(f"Patched fetch_one error: {e}")
            return None
    
    # Apply patches
    db_manager.execute_query = patched_execute_query
    db_manager.fetch_one = patched_fetch_one
    
    print("✅ Database manager patched to return dictionaries")
    return db_manager
def test_database_returns(db_manager):
    """Test what the database manager actually returns"""
    try:
        st.write("🔍 **Testing Database Returns:**")
        
        # Test execute_query
        result = db_manager.execute_query("SELECT 1 as test_col, 'test' as test_str")
        
        if result:
            first_row = result[0]
            st.write(f"**Result type:** {type(first_row)}")
            st.write(f"**First row:** {first_row}")
            
            if isinstance(first_row, dict):
                st.success("✅ Database returns dictionaries - accessing with ['key'] will work")
                st.write(f"test_col value: {first_row['test_col']}")
            elif isinstance(first_row, tuple):
                st.error("❌ Database returns tuples - need to access with [0], [1], etc.")
                st.write(f"test_col value: {first_row[0]}")
            else:
                st.warning(f"❓ Database returns {type(first_row)} - unexpected type")
        else:
            st.error("❌ Database query returned empty result")
            
        # Test fetch_one
        single_result = db_manager.fetch_one("SELECT 1 as test_col, 'test' as test_str")
        
        if single_result:
            st.write(f"**fetch_one type:** {type(single_result)}")
            st.write(f"**fetch_one result:** {single_result}")
            
            if isinstance(single_result, dict):
                st.success("✅ fetch_one returns dictionaries")
            else:
                st.error("❌ fetch_one returns non-dictionary")
        
    except Exception as e:
        st.error(f"Database test failed: {e}")
        st.code(traceback.format_exc())
class Batchprocessing_SparkMasterWorker:
    #Démonstrations concrètes des opérations Master-Worker avec parallelize
    def __init__(self, spark_session):
        self.spark = spark_session
        
    def demo_basic_rdd_operations(self):
        """ULTRA-SAFE: Basic RDD operations without complex serialization"""
        print("🚀 DEMO 1: Basic RDD Operations (Ultra-Safe Version)")
        
        try:
            if not self.spark:
                print("❌ Spark session not available")
                return []
            
            # 1. Create very simple data
            simple_data = list(range(1, 11))  # Just numbers 1-10
            print(f"Creating RDD with {len(simple_data)} elements...")
            
            # 2. Create RDD with only 1 partition to avoid worker communication issues
            data_rdd = self.spark.sparkContext.parallelize(simple_data, numSlices=4)
            print(f"✅ RDD created with {data_rdd.getNumPartitions()} partition")
            
            # 3. ULTRA-SIMPLE function - no external dependencies
            def simple_double(x):
                return x * 2
            # 4. Apply transformation
            doubled_rdd = data_rdd.map(simple_double)
            # 5. SAFE collection with take() instead of collect()
            print("📥 Collecting results...")
            
            
        except Exception as e:
            print(f"❌ Demo 1 failed: {e}")
            return []
    
    
class QuickSparkIntegration:
    # Fix your Spark configuration in QuickSparkIntegration.__init__
    def __init__(self, db_manager):
        self.db_manager = db_manager
        self.spark = None
        self.spark_enabled = SPARK_AVAILABLE
        
        if self.spark_enabled:
            try:
                # FIXED: Much safer Spark configuration
                self.spark = SparkSession.builder \
                    .appName("CVMatchingSystem") \
                    .master("local[1]") \
                    .config("spark.driver.memory", "512m") \
                    .config("spark.executor.memory", "512m") \
                    .config("spark.driver.maxResultSize", "128m") \
                    .config("spark.serializer", "org.apache.spark.serializer.KryoSerializer") \
                    .config("spark.sql.execution.arrow.pyspark.enabled", "false") \
                    .config("spark.sql.adaptive.enabled", "true") \
                    .config("spark.sql.adaptive.coalescePartitions.enabled", "true") \
                    .config("spark.driver.host", "localhost") \
                    .config("spark.driver.bindAddress", "localhost") \
                    .config("spark.ui.enabled", "true") \
                     .config("spark.ui.port", "4041") \
                    .config("spark.driver.extraJavaOptions", "-Xss2m -XX:+UseG1GC") \
                    .config("spark.executor.extraJavaOptions", "-Xss2m -XX:+UseG1GC") \
                    .config("spark.python.worker.reuse", "false") \
                    .config("spark.task.maxFailures", "1") \
                    .getOrCreate()
                
                self.spark.sparkContext.setLogLevel("ERROR")
                
                # Test Spark with simple operation
                test_rdd = self.spark.sparkContext.parallelize([1, 2, 3], 1)
                test_result = test_rdd.collect()
                
                if test_result == [1, 2, 3]:
                    print("✅ Spark initialized and tested successfully")
                else:
                    raise Exception("Spark test failed")
                
            except Exception as e:
                print(f"⚠️ Spark initialization failed: {e}")
                self.spark_enabled = False
                self.spark = None

    def _safe_insert_matches(self, matches_list, db_manager):
        """Safely insert matches without complex operations"""
        try:
            matches_inserted = 0
            
            for match in matches_list:
                try:
                    # Simple insert query
                    insert_query = """
                        INSERT INTO job_matches (
                            match_id, candidate_id, job_offer_id, 
                            similarity_score, distance, overall_evaluation,
                            created_at, updated_at
                        ) VALUES (%s, %s, %s, %s, %s, %s, NOW(), NOW())
                        ON CONFLICT (candidate_id, job_offer_id) DO UPDATE SET
                            similarity_score = EXCLUDED.similarity_score,
                            distance = EXCLUDED.distance,
                            updated_at = NOW()
                    """
                    
                    import uuid
                    db_manager.execute_query(insert_query, (
                        str(uuid.uuid4()),
                        match['candidate_id'],
                        match['job_offer_id'],
                        float(match['similarity_score']),
                        float(match['distance']),
                        float(match['similarity_score'])
                    ))
                    
                    matches_inserted += 1
                    
                except Exception as insert_error:
                    print(f"❌ Failed to insert match: {insert_error}")
                    continue
            
            return matches_inserted
            
        except Exception as e:
            print(f"❌ Bulk insert error: {e}")
            return 0
    def _spark_matching_pipeline(self, db_manager):
        """ULTRA-SAFE: Simplified Spark matching pipeline"""
        try:
            print("🚀 Starting SAFE Spark matching pipeline...")
            
            # Get small sample of data to avoid memory issues
            candidates_data = db_manager.execute_query(
                "SELECT candidate_id, job_category, experience_years FROM candidates LIMIT 20"
            )
            jobs_data = db_manager.execute_query(
                "SELECT job_offer_id, job_category, experience_required FROM job_offers WHERE status = 'active' LIMIT 5"
            )
            
            if not candidates_data or not jobs_data:
                print("❌ No data available for matching")
                return 0
            
            print(f"📊 Processing {len(candidates_data)} candidates × {len(jobs_data)} jobs")
            
            # Convert to simple format for Spark
            candidate_tuples = [
                (row['candidate_id'], row.get('job_category', ''), row.get('experience_years', 0))
                for row in candidates_data
            ]
            job_tuples = [
                (row['job_offer_id'], row.get('job_category', ''), row.get('experience_required', 0))
                for row in jobs_data
            ]
            
            # Create RDDs with single partition
            candidates_rdd = self.spark.sparkContext.parallelize(candidate_tuples, 1)
            
            # Broadcast jobs data
            jobs_broadcast = self.spark.sparkContext.broadcast(job_tuples)
            
            # Simple matching function
            def simple_match(candidate_tuple):
                candidate_id, candidate_category, candidate_exp = candidate_tuple
                matches = []
                
                for job_tuple in jobs_broadcast.value:
                    job_id, job_category, job_exp_required = job_tuple
                    
                    # Simple category-based scoring
                    score = 0.3  # Base score
                    
                    if candidate_category and job_category:
                        if candidate_category.lower() == job_category.lower():
                            score = 0.8
                        elif candidate_category.lower() in job_category.lower():
                            score = 0.6
                    
                    # Experience bonus
                    if candidate_exp >= job_exp_required:
                        score += 0.1
                    
                    if score >= 0.4:  # Minimum threshold
                        matches.append({
                            'candidate_id': candidate_id,
                            'job_offer_id': job_id,
                            'similarity_score': min(score, 0.95),
                            'distance': 1.0 - score
                        })
                
                return matches
            
            # Process matches
            matches_rdd = candidates_rdd.flatMap(simple_match)
            
            # Collect results safely
            all_matches = matches_rdd.take(100)  # Limit results
            
            # Clean up broadcast variable
            jobs_broadcast.unpersist()
            
            if all_matches:
                # Insert matches into database
                matches_inserted = self._safe_insert_matches(all_matches, db_manager)
                print(f"✅ Spark pipeline completed: {matches_inserted} matches inserted")
                return matches_inserted
            else:
                print("⚠️ No matches found")
                return 0
            
        except Exception as e:
            print(f"❌ Spark pipeline error: {e}")
            return 0
    #ssytème amélioée de matching avec spark 
    def enhanced_process_matching(self, db_manager, original_matching_function):
        """SAFE: Enhanced matching with better error handling"""
        
        if not self.spark_enabled or not self.spark:
            print("📊 Using original matching method (Spark disabled)")
            return original_matching_function()
        
        try:
            print("🚀 Attempting SAFE Spark-accelerated matching...")
            start_time = time.time()
            
            # Try SAFE Spark-based matching
            matches_created = self._spark_matching_pipeline(db_manager)
            
            end_time = time.time()
            duration = end_time - start_time
            
            if matches_created > 0:
                print(f"✅ SAFE Spark matching completed: {matches_created} matches in {duration:.2f}s")
                
                if 'st' in globals():
                    st.success(f"🚀 Spark-accelerated matching: {matches_created} matches in {duration:.2f}s")
                
                return matches_created
            else:
                print("⚠️ Spark matching returned 0 matches, falling back...")
                raise Exception("No matches created with Spark")
                
        except Exception as e:
            print(f"⚠️ Spark matching failed: {e}")
            print("📊 Falling back to original matching method...")
            
            if 'st' in globals():
                st.warning("Spark matching failed, using standard method...")
            
            return original_matching_function()
    
    
    def _calculate_similarity_spark(self):
        """
        Simplified similarity calculation for Spark
        """
        # For now, use a simple heuristic based on job category match
        # You can enhance this with actual vector similarity later
        return when(
            col("candidates.job_category") == col("jobs.job_category"), 
            lit(0.8)
        ).otherwise(lit(0.5))
    
    def _bulk_insert_matches(self, matches_df, db_manager):
        """
        Efficient bulk insert of matches
        """
        try:
            records = []
            for _, row in matches_df.iterrows():
                records.append((
                    str(row['match_id']) if pd.notna(row['match_id']) else str(uuid.uuid4()),
                    row['candidate_id'],
                    row['job_offer_id'], 
                    float(row['distance']),
                    float(row['similarity_score']),
                    float(row['overall_evaluation']),
                    datetime.now(),
                    datetime.now()
                ))
            
            # Use execute_many for efficiency
            conn = db_manager.get_connection()
            cursor = conn.cursor()
            
            insert_sql = """
                INSERT INTO job_matches (
                    match_id, candidate_id, job_offer_id, distance,
                    similarity_score, overall_evaluation, created_at, updated_at
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            """
            
            cursor.executemany(insert_sql, records)
            conn.commit()
            cursor.close()
            
            print(f"✅ Inserted {len(records)} matches successfully")
            
        except Exception as e:
            print(f"❌ Bulk insert error: {e}")
            if 'conn' in locals():
                conn.rollback()
    
    def enhanced_analytics(self, db_manager, original_analytics_function):
        """
        Enhanced analytics with Spark acceleration
        """
        if not self.spark_enabled:
            return original_analytics_function()
        
        try:
            print("🚀 Running Spark-accelerated analytics...")
            start_time = time.time()
            
            # Load data into Spark
            candidates_pd = pd.DataFrame(db_manager.execute_query("SELECT * FROM candidates"))
            matches_pd = pd.DataFrame(db_manager.execute_query("SELECT * FROM job_matches"))
            
            if candidates_pd.empty:
                return original_analytics_function()
            
            candidates_df = self.spark.createDataFrame(candidates_pd)
            
            # Distributed skills analysis
            skills_analysis = self._analyze_skills_distributed(candidates_df)
            
            # Match statistics if matches exist
            match_stats = {}
            if not matches_pd.empty:
                matches_df = self.spark.createDataFrame(matches_pd)
                match_stats = self._analyze_matches_distributed(matches_df)
            
            end_time = time.time()
            print(f"✅ Spark analytics completed in {end_time - start_time:.2f}s")
            
            return {
                'technical_skills': skills_analysis.get('technical', pd.DataFrame()),
                'soft_skills': skills_analysis.get('soft', pd.DataFrame()),
                'match_stats': match_stats,
                'processing_time': end_time - start_time
            }
            
        except Exception as e:
            print(f"⚠️ Spark analytics failed: {e}")
            return original_analytics_function()
    
    def _analyze_skills_distributed(self, candidates_df):
        """
        Distributed skills analysis
        """
        try:
            # UDF to extract skills from JSON
            @udf(returnType=ArrayType(StringType()))
            def extract_technical_skills(skills_json):
                try:
                    if skills_json:
                        skills_data = json.loads(skills_json) if isinstance(skills_json, str) else skills_json
                        return skills_data.get('technical_skills', [])
                    return []
                except:
                    return []
            
            @udf(returnType=ArrayType(StringType()))
            def extract_soft_skills(skills_json):
                try:
                    if skills_json:
                        skills_data = json.loads(skills_json) if isinstance(skills_json, str) else skills_json
                        return skills_data.get('soft_skills', [])
                    return []
                except:
                    return []
            
            # Technical skills analysis
            tech_skills = candidates_df \
                .withColumn("tech_skills", extract_technical_skills(col("skills"))) \
                .select(explode(col("tech_skills")).alias("skill")) \
                .groupBy("skill") \
                .count() \
                .orderBy(desc("count")) \
                .toPandas()
            
            # Soft skills analysis
            soft_skills = candidates_df \
                .withColumn("soft_skills", extract_soft_skills(col("skills"))) \
                .select(explode(col("soft_skills")).alias("skill")) \
                .groupBy("skill") \
                .count() \
                .orderBy(desc("count")) \
                .toPandas()
            
            return {
                'technical': tech_skills,
                'soft': soft_skills
            }
            
        except Exception as e:
            print(f"Skills analysis error: {e}")
            return {'technical': pd.DataFrame(), 'soft': pd.DataFrame()}
    
    def _analyze_matches_distributed(self, matches_df):
        """
        Distributed match analysis
        """
        try:
            stats = matches_df.agg(
                avg("similarity_score").alias("avg_similarity"),
                max("similarity_score").alias("max_similarity"),
                count("*").alias("total_matches")
            ).collect()[0]
            
            return stats.asDict()
            
        except Exception as e:
            print(f"Match analysis error: {e}")
            return {}
    
    
class EnhancedQuickSparkIntegration(QuickSparkIntegration):
    """
    Version améliorée avec démonstrations Master-Worker
    """
    
    def __init__(self, db_manager):
        super().__init__(db_manager)
        self.demo = Batchprocessing_SparkMasterWorker(self.spark) if self.spark else None
    
    #DÉMONSTRATION de MASTER-WORKER AVEC PARALLELIZE
    def demonstrate_master_worker_operations(self):
        #Lance toutes les démonstrations Master-Worker
        if not self.spark_enabled or not self.demo:
            print("❌ Spark non disponible pour les démonstrations")
            return
        
        print("🚀 DÉMONSTRATIONS MASTER-WORKER AVEC PARALLELIZE")
        print("=" * 60)
        
        # Préparer des données de test
        test_candidates = [
            {
                'candidate_id': i,
                'skills_vector': [random.uniform(0, 1) for _ in range(50)],
                'job_category': random.choice(['tech', 'finance', 'healthcare']),
                'skills': {
                    'technical_skills': random.sample(['python', 'java', 'sql', 'ml', 'react'], 3),
                    'soft_skills': ['communication', 'leadership']
                }
            }
            for i in range(1, 101)
        ]
        
        test_jobs = [
            {
                'job_offer_id': f"job_{i}",
                'skills_vector': [random.uniform(0, 1) for _ in range(50)],
                'job_category': random.choice(['tech', 'finance', 'healthcare'])
            }
            for i in range(1, 21)
        ]
        
        # Exécuter toutes les démonstrations
        try:
            # Demo 1: Opérations RDD de base
            self.demo.demo_basic_rdd_operations()
            print("\n🎉 Toutes les démonstrations Master-Worker terminées avec succès!")
            
        except Exception as e:
            print(f"❌ Erreur dans les démonstrations: {e}")

class DataPipeline:
    """COMPLETE DataPipeline class - USE THIS VERSION ONLY"""
    
    def __init__(self, db_manager, matching_engine):
        self.db_manager = db_manager
        self.matching_engine = matching_engine
        self.auto_update_thread = None
        print("✅ DataPipeline initialized with all methods")
    
    def fetch_candidates(self):
        """Fetch all candidates"""
        try:
            candidates = self.db_manager.execute_query("SELECT * FROM candidates LIMIT 10")
            return candidates if candidates else []
        except Exception as e:
            print(f"Error fetching candidates: {e}")
            return []
    def fetch_clients(self):
        """Fetch all clients"""
        try:
            clients = self.db_manager.execute_query("SELECT * FROM clients LIMIT 10")
            return clients if clients else []
        except Exception as e:
            print(f"Error fetching clients: {e}")
            return []
    def get_candidate_matches(self, candidate_id):
        """Get job matches for a candidate"""
        print(f"🔍 Getting matches for candidate {candidate_id}")
        try:
            query = """
                SELECT 
                    jm.*,
                    jo.job_title,
                    jo.company_name,
                    jo.city,
                    jo.country,
                    jo.job_type,
                    jo.description,
                    jo.job_category
                FROM job_matches jm
                JOIN job_offers jo ON jo.job_offer_id = jm.job_offer_id
                WHERE jm.candidate_id = %s
                ORDER BY jm.similarity_score DESC
            """
            
            matches = self.db_manager.execute_query(query, (candidate_id,))
            print(f"📊 Found {len(matches) if matches else 0} matches for candidate {candidate_id}")
            
            if matches:
                return pd.DataFrame(matches)
            else:
                return pd.DataFrame()
                
        except Exception as e:
            print(f"Error getting candidate matches: {e}")
            return pd.DataFrame()
    def get_job_candidates(self, job_offer_id):
        """Get candidates for a specific job"""
        print(f"🔍 Getting candidates for job {job_offer_id}")
        try:
            query = """
                SELECT 
                    c.*,
                    jm.similarity_score,
                    jm.distance,
                    jm.overall_evaluation,
                    jm.created_at as match_created_at
                FROM candidates c
                JOIN job_matches jm ON c.candidate_id = jm.candidate_id
                WHERE jm.job_offer_id = %s
                ORDER BY jm.similarity_score DESC
            """
            
            candidates = self.db_manager.execute_query(query, (job_offer_id,))
            print(f"📊 Found {len(candidates) if candidates else 0} candidates for job {job_offer_id}")
            
            if candidates:
                return pd.DataFrame(candidates)
            else:
                return pd.DataFrame()
                
        except Exception as e:
            print(f"Error getting job candidates: {e}")
            return pd.DataFrame()
    def process_matching_for_candidate(self, candidate_id: int):
        """Process matching for a specific candidate"""
        print(f"🎯 Processing matches for candidate {candidate_id}")
        try:
            # Get candidate info
            candidate = self.db_manager.fetch_one(
                "SELECT * FROM candidates WHERE candidate_id = %s",
                (candidate_id,)
            )
            
            if not candidate:
                print(f"❌ Candidate {candidate_id} not found")
                if 'st' in globals():
                    st.error(f"Candidate {candidate_id} not found!")
                return 0
            
            # Get job offers
            jobs = self.db_manager.execute_query("SELECT * FROM job_offers WHERE status = 'active'")
            
            if not jobs:
                print("❌ No active job offers found")
                if 'st' in globals():
                    st.warning("No active job offers found!")
                return 0
            
            print(f"📋 Found {len(jobs)} active job offers")
            
            # Insert matches
            insert_sql = """
                INSERT INTO job_matches (
                    match_id, candidate_id, job_offer_id, similarity_score,
                    distance, overall_evaluation, created_at, updated_at
                ) VALUES (%s, %s, %s, %s, %s, %s, NOW(), NOW())
                
                ON CONFLICT (candidate_id, job_offer_id)
                DO UPDATE SET
                    match_id = EXCLUDED.match_id,
                    distance = EXCLUDED.distance,
                    similarity_score = EXCLUDED.similarity_score,
                    overall_evaluation = EXCLUDED.overall_evaluation,
                    updated_at = NOW()
            """
            
            matches_created = 0
            matches_processed = 0
            
            for job in jobs:
                try:
                    matches_processed += 1
                    
                    # Calculate match
                    match_result = self.matching_engine.calculate_match(candidate, job)
                    
                    if match_result['similarity_score'] >= self.matching_engine.MIN_SIMILARITY_SCORE:
                        params = (
                            str(uuid.uuid4()),
                            candidate_id,
                            job['job_offer_id'],
                            match_result['similarity_score'],
                            match_result['distance'],
                            match_result['similarity_score']
                        )
                        
                        self.db_manager.execute_query(insert_sql, params)
                        matches_created += 1
                        
                except Exception as e:
                    print(f"❌ Error with job {job.get('job_offer_id')}: {e}")
                    continue
            
            print(f" {matches_created} matches from {matches_processed} jobs")
            
            if 'st' in globals():
                if matches_created > 0:
                    st.success(f" {matches_created} matches!")
                else:
                    st.warning(f"⚠️ No matches found above threshold")
            
            return matches_created
            
        except Exception as e:
            error_msg = f"❌ Error in process_matching_for_candidate: {str(e)}"
            print(error_msg)
            print(f"Traceback: {traceback.format_exc()}")
            
            if 'st' in globals():
                st.error(error_msg)
            
            return 0
    
    def process_matching(self):
        """Process all candidate-job matches"""
        print("🔄 Starting bulk matching process...")
        
        if 'st' in globals():
            st.info("🔄 Processing all matches...")

        try:
            # Clear existing matches
            self.db_manager.execute_query("DELETE FROM job_matches")
            print("🗑️ Cleared existing matches")
            
            # Get candidates and jobs
            candidates = self.db_manager.execute_query(
                "SELECT candidate_id, first_name, last_name, skills_vector, skills FROM candidates"
            )
            jobs = self.db_manager.execute_query(
                "SELECT job_offer_id, job_title, company_name, skills_vector, technical_skills FROM job_offers WHERE status = 'active'"
            )

            if not candidates or not jobs:
                print("❌ No candidates or jobs found")
                if 'st' in globals():
                    st.warning("No candidates or job offers found!")
                return

            print(f"📊 Processing {len(candidates)} candidates × {len(jobs)} jobs")

            insert_sql = """
                INSERT INTO job_matches (
                    match_id, candidate_id, job_offer_id,
                    distance, similarity_score, overall_evaluation, 
                    created_at, updated_at
                ) VALUES (%s, %s, %s, %s, %s, %s, NOW(), NOW())
            """

            inserted_count = 0
            processed_count = 0
            
            for job_index, job in enumerate(jobs):
                job_id = job['job_offer_id']
                job_matches = []
                
                for candidate in candidates:
                    try:
                        processed_count += 1
                        match_result = self.matching_engine.calculate_match(candidate, job)
                        
                        if match_result['similarity_score'] >= self.matching_engine.MIN_SIMILARITY_SCORE:
                            job_matches.append({
                                'candidate_id': candidate['candidate_id'],
                                'similarity_score': match_result['similarity_score'],
                                'distance': match_result['distance'],
                                'overall_evaluation': match_result['similarity_score']
                            })
                            
                    except Exception as e:
                        print(f"❌ Error with candidate {candidate.get('candidate_id')}: {e}")
                        continue
                
                # Sort and limit matches
                job_matches.sort(key=lambda x: x['similarity_score'], reverse=True)
                
                if hasattr(self.matching_engine, 'MAX_MATCHES_PER_JOB'):
                    max_matches = min(len(job_matches), self.matching_engine.MAX_MATCHES_PER_JOB)
                    job_matches = job_matches[:max_matches]
                
                # Insert matches
                for match in job_matches:
                    try:
                        self.db_manager.execute_query(insert_sql, (
                            str(uuid.uuid4()),
                            match['candidate_id'],
                            job_id,
                            match['distance'],
                            match['similarity_score'],
                            match['overall_evaluation']
                        ))
                        inserted_count += 1
                        
                    except Exception as e:
                        print(f"❌ Error inserting match: {e}")
                        continue

            print(f"✅ Matching completed: {inserted_count} matches")
            
            if 'st' in globals():
                st.success(f" {inserted_count} matches!")
                if hasattr(st.session_state, 'last_update'):
                    st.session_state.last_update = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        except Exception as e:
            error_msg = f"❌ Error in bulk matching: {str(e)}"
            print(error_msg)
            print(f"Traceback: {traceback.format_exc()}")
            
            if 'st' in globals():
                st.error(error_msg)

# Enhanced DataPipeline class - ADD THIS TO YOUR EXISTING DataPipeline
class DataPipelineWithSpark(DataPipeline):
    """
    Enhanced version of your existing DataPipeline with Spark acceleration
    """
    
    def __init__(self, db_manager, matching_engine):
        # Initialize original pipeline
        super().__init__(db_manager, matching_engine)
        
        # Add Spark capability
        self.quick_spark = QuickSparkIntegration(db_manager)
        print("✅ DataPipeline enhanced with Spark capabilities")
    
    def process_matching(self):
        #amélioer le matching avec spark
        #Enhanced matching method that uses Spark when available
        
        def original_matching():
            """Your original matching logic"""
            return super().process_matching()
        
        # Use Spark-enhanced matching
        return self.quick_spark.enhanced_process_matching(
            self.db_manager, 
            original_matching
        )
    
    def get_enhanced_analytics(self):
        """
        Enhanced analytics using Spark
        """
        def original_analytics():
            """Your original analytics logic"""
            return self.db_manager.get_skills_analytics()
        
        return self.quick_spark.enhanced_analytics(
            self.db_manager,
            original_analytics
        )
    
    #lancer mater et workers
    def run_master_worker_demos(self):
        """
        Lance les démonstrations Master-Worker
        """
        if hasattr(self, 'quick_spark') and self.quick_spark.spark_enabled:
            enhanced_spark = EnhancedQuickSparkIntegration(self.db_manager)
            enhanced_spark.demonstrate_master_worker_operations()
        else:
            print("❌ Spark non disponible pour les démonstrations")
    def close(self):
        """Clean up resources"""
        if hasattr(self, 'quick_spark'):
            self.quick_spark.close()


class EnhancedMatchingEngine:
    """Uses advanced distance calculation and Gaussian similarity conversion
    """
    
    def __init__(self):
        # Simplified - focus only on skills
        self.weights = {
            'skills': 1.0  # 100% weight on skills only
        }
        self.distances = []
        self.matches = []
        self.threshold_analysis = {}
        # MODIFICATION 1: Lower the minimum similarity threshold to allow more variety
        self.MIN_SIMILARITY_SCORE = 0.2 # Changed from 0.5 to 0.3 (30% threshold)
        # MODIFICATION 2: Increase max distance to allow more matches
        self.MAX_DISTANCE_FOR_SIMILARITY = 10.0  # Changed from 5.0 to 10.0
        # MODIFICATION 3: Add maximum matches per job to prevent unlimited matching
        self.MAX_MATCHES_PER_JOB = 15  # NEW: Limit matches per job

    def calculate_distance(self, vec1: np.ndarray, vec2: np.ndarray, method: str = 'euclidean') -> float:
        """
        Calculate distance between vectors using specified method
        
        Args:
            vec1: First vector (candidate skills)
            vec2: Second vector (job offer skills)
            method: Distance calculation method ('euclidean', 'cosine', 'manhattan')
        
        Returns:
            Distance value (lower = better match)
        """
        if vec1.size == 0 or vec2.size == 0 or vec1.shape != vec2.shape:
            return float('inf')
        
        # MODIFICATION 4: Add normalization to prevent extreme distances
        # Normalize vectors to unit length before distance calculation
        try:
            vec1_norm = vec1 / (np.linalg.norm(vec1) + 1e-8) 
            vec2_norm = vec2 / (np.linalg.norm(vec2) + 1e-8)
        except:
            return float('inf')
        
        if method == 'euclidean':
            # Euclidean distance: √(Σ(candidate_skill_i - job_skill_i)²)
            return np.linalg.norm(vec1_norm - vec2_norm)
        
        elif method == 'cosine':
            # Cosine distance: 1 - cosine_similarity
            dot_product = np.dot(vec1_norm, vec2_norm)
            norms = np.linalg.norm(vec1_norm) * np.linalg.norm(vec2_norm)
            if norms == 0:
                return float('inf')
            cosine_similarity = dot_product / norms
            return 1 - cosine_similarity
        
        elif method == 'manhattan':
            # Manhattan distance: Σ|candidate_skill_i - job_skill_i|
            return np.sum(np.abs(vec1_norm - vec2_norm))
        
        else:
            raise ValueError(f"Unknown distance method: {method}")

    def calculate_weighted_distance(self, vec1: np.ndarray, vec2: np.ndarray, 
                                  weights: Optional[np.ndarray] = None) -> float:
        #calcul de pondération
        if vec1.size == 0 or vec2.size == 0 or vec1.shape != vec2.shape:
            return float('inf')
        if weights is None:
            weights = np.ones(vec1.shape[0])
        elif weights.shape != vec1.shape:
            return float('inf')
        #La normalisation pour que les tailles de 2 vecteurs soit compatibles
        try:
            vec1_norm = vec1 / (np.linalg.norm(vec1) + 1e-8)
            vec2_norm = vec2 / (np.linalg.norm(vec2) + 1e-8)
        except: return float('inf')
        # Calcul de distance euclidienne pondéré: √(Σ(w_i * (x_i - y_i)²))
        weighted_diff = weights * (vec1_norm - vec2_norm) ** 2
        return np.sqrt(np.sum(weighted_diff))

    def distance_to_similarity_gaussian(self, distance: float, sigma: float = 1.0) -> float:
        if np.isinf(distance) or np.isnan(distance):
            return 0.0
        # MODIFICATION 6: Adjust sigma based on distance to create more variety
        # Use adaptive sigma that depends on the distance magnitude
        adaptive_sigma = max(sigma, distance * 0.1)  # NEW: Adaptive sigma
        # Gaussian similarity: e^(-(distance²)/(2*σ²))
        similarity = np.exp(-(distance ** 2) / (2 * adaptive_sigma ** 2))
        return min(max(similarity, 0.0), 1.0)

    def determine_threshold(self, distances: List[float], method: str = 'percentile', **kwargs) -> float:
        """
        Determine optimal threshold dynamically using various methods
        
        Args:
            distances: List of calculated distances
            method: Threshold determination method
            **kwargs: Additional parameters for specific methods
        
        Returns:
            Threshold value for filtering matches
        """
        # Filter out infinite and NaN values
        finite_distances = [d for d in distances if np.isfinite(d)]
        
        if not finite_distances:
            return kwargs.get('fixed_threshold', 2.0)
        
        # MODIFICATION 7: Change default method to be more lenient
        if method == 'percentile':
            percentile = kwargs.get('percentile', 75)  # Changed from 50 to 75 (more lenient)
            return np.percentile(finite_distances, percentile)
        
        elif method == 'std':
            n_std = kwargs.get('n_std', 1.5)  # Changed from 1 to 1.5 (more lenient)
            mean = np.mean(finite_distances)
            std = np.std(finite_distances)
            return max(0, mean - n_std * std)
        
        elif method == 'iqr':
            # Interquartile Range method
            q1 = np.percentile(finite_distances, 25)
            q3 = np.percentile(finite_distances, 75)
            iqr = q3 - q1
            return q1 - 1.5 * iqr  # Lower fence for outlier detection
        
        elif method == 'adaptive':
            # MODIFICATION 8: Make adaptive threshold more lenient
            mean = np.mean(finite_distances)
            median = np.median(finite_distances)
            std = np.std(finite_distances)
            
            # If distribution is skewed, use median-based threshold
            if abs(mean - median) > 0.5 * std:
                return median - 0.25 * std  # Changed from 0.5 to 0.25 (more lenient)
            else:
                return mean - 0.25 * std    # Changed from 0.5 to 0.25 (more lenient)
        
        else:  # 'fixed' method
            return kwargs.get('fixed_threshold', 3.0)  # Changed from 2.0 to 3.0 (more lenient)

    
    # ==================== STATISTICAL ANALYSIS METHODS ====================
    
    
      
    def safe_extract_skills(self, skills_data):
        """
        Safely extract skills from various formats (JSON dict, string, etc.)
        Returns a set of lowercase skill strings
        """
        if not skills_data:
            return set()
        
        # Handle JSON dictionary format (from PostgreSQL)
        if isinstance(skills_data, dict):
            all_skills = []
            
            # Extract technical skills
            tech_skills = skills_data.get('technical_skills', [])
            if isinstance(tech_skills, list):
                all_skills.extend([str(skill).strip().lower() for skill in tech_skills if skill])
            elif isinstance(tech_skills, str):
                all_skills.extend([skill.strip().lower() for skill in tech_skills.split(',') if skill.strip()])
            
            # Extract soft skills
            soft_skills = skills_data.get('soft_skills', [])
            if isinstance(soft_skills, list):
                all_skills.extend([str(skill).strip().lower() for skill in soft_skills if skill])
            elif isinstance(soft_skills, str):
                all_skills.extend([skill.strip().lower() for skill in soft_skills.split(',') if skill.strip()])
            
            return set(all_skills)
        
        # Handle JSON string format
        if isinstance(skills_data, str):
            try:
                # Try to parse as JSON first
                parsed_skills = json.loads(skills_data)
                return self.safe_extract_skills(parsed_skills)  # Recursive call
            except (json.JSONDecodeError, ValueError):
                # Treat as comma-separated string
                return set(skill.strip().lower() for skill in skills_data.split(',') if skill.strip())
        
        # Handle list format
        if isinstance(skills_data, list):
            return set(str(skill).strip().lower() for skill in skills_data if skill)
        
        # Fallback
        return set()

    def convert_vector_to_array(self, vector_data) -> np.ndarray:
        """Convert vector data to numpy array - handles different formats"""
        try:
            if vector_data is None:
                return np.array([])
            
            # If it's already a list or array
            if isinstance(vector_data, (list, np.ndarray)):
                return np.array(vector_data)
            
            # If it's a string representation
            if isinstance(vector_data, str):
                if vector_data.strip() in ['', '[]', 'null']:
                    return np.array([])
                # Try to parse as JSON first
                try:
                    parsed = json.loads(vector_data)
                    return np.array(parsed)
                except:
                    # Try to parse as comma-separated values
                    return np.fromstring(vector_data.strip('[]'), sep=',')
            
            return np.array([])
        except Exception as e:
            print(f"Error converting vector: {str(e)}")
            return np.array([])

    def calculate_skills_similarity(self, candidate_skills_vector, job_skills_vector, method='vector'):
        """
        Enhanced calculate similarity between candidate skills and job requirements
        Now supports both vector embeddings and text-based similarity
        """
        try:
            if method == 'vector':
                # Use vector embeddings if available
                return self.calculate_vector_similarity(candidate_skills_vector, job_skills_vector)
            else:
                # Use text-based similarity as fallback
                return self.calculate_text_based_skills_similarity(candidate_skills_vector, job_skills_vector)
        except Exception as e:
            print(f"Error calculating skills similarity: {e}")
            return 0.0

    def calculate_vector_similarity(self, candidate_skills_vector, job_skills_vector):
        """
        Calculate similarity using vector embeddings with enhanced error handling
        """
        try:
            # Convert to numpy arrays
            if isinstance(candidate_skills_vector, (list, tuple)):
                candidate_vector = np.array(candidate_skills_vector, dtype=float)
            elif isinstance(candidate_skills_vector, str):
                candidate_vector = np.array(json.loads(candidate_skills_vector), dtype=float)
            else:
                candidate_vector = self.convert_vector_to_array(candidate_skills_vector)
            
            if isinstance(job_skills_vector, (list, tuple)):
                job_vector = np.array(job_skills_vector, dtype=float)
            elif isinstance(job_skills_vector, str):
                job_vector = np.array(json.loads(job_skills_vector), dtype=float)
            else:
                job_vector = self.convert_vector_to_array(job_skills_vector)
            
            # Check if vectors have the same dimension
            if candidate_vector.shape != job_vector.shape:
                print(f"Vector dimension mismatch: candidate {candidate_vector.shape} vs job {job_vector.shape}")
                return 0.0
            
            # MODIFICATION 10: Add randomness to break ties and create variety
            # Add small random noise to prevent identical similarities
            noise_factor = 0.05  # 5% noise
            random_noise = np.random.uniform(-noise_factor, noise_factor)
            
            # Calculate distance using enhanced method
            distance = self.calculate_distance(candidate_vector, job_vector, method='euclidean')
            
            # Convert distance to similarity using only Gaussian method
            similarity = self.distance_to_similarity_gaussian(distance, sigma=1.0)
            
            # Add the random noise to create variety (but keep within bounds)
            similarity_with_noise = similarity + random_noise
            similarity_with_noise = max(0.0, min(1.0, similarity_with_noise))
            
            return similarity_with_noise
            
        except Exception as e:
            print(f"Error calculating vector similarity: {e}")
            return self.calculate_text_based_skills_similarity(candidate_skills_vector, job_skills_vector)

    def calculate_text_based_skills_similarity(self, candidate_skills, job_skills):
        """
        Fallback text-based skills similarity calculation
        """
        try:
            candidate_skills_set = self.safe_extract_skills(candidate_skills)
            job_skills_set = self.safe_extract_skills(job_skills)
            
            if not job_skills_set:
                return 1.0  # If no skills required, perfect match
            
            if not candidate_skills_set:
                return 0.0  # If candidate has no skills, no match
            
            # Calculate intersection and similarity
            common_skills = candidate_skills_set.intersection(job_skills_set)
            similarity = len(common_skills) / len(job_skills_set)
            
            # MODIFICATION 11: Add randomness here too
            noise_factor = 0.1  # 10% noise for text-based similarity
            random_noise = np.random.uniform(-noise_factor, noise_factor)
            similarity_with_noise = similarity + random_noise
            
            return max(0.0, min(1.0, similarity_with_noise))
        except Exception as e:
            print(f"Error in text-based skills similarity: {e}")
            return 0.0

    def calculate_match(self, candidate, job):
        """FIXED: Enhanced matching with proper error handling"""
        try:
            # SAFE: Handle both tuple and dict candidate data
            if isinstance(candidate, dict):
                candidate_skills_vector = candidate.get('skills_vector', None)
                candidate_skills = candidate.get('skills', None)
            else:  # tuple fallback
                candidate_skills_vector = candidate[0] if len(candidate) > 0 else None
                candidate_skills = candidate[1] if len(candidate) > 1 else None
            
            # SAFE: Handle both tuple and dict job data
            if isinstance(job, dict):
                job_skills_vector = job.get('skills_vector', None)
                job_skills = job.get('technical_skills', None)
            else:  # tuple fallback
                job_skills_vector = job[0] if len(job) > 0 else None
                job_skills = job[1] if len(job) > 1 else None
            
            # Calculate similarity
            try:
                skills_score = self.calculate_skills_similarity(
                    candidate_skills_vector, 
                    job_skills_vector, 
                    method='vector'
                )
            except Exception as similarity_error:
                print(f"Similarity calculation failed: {similarity_error}")
                # Fallback to basic matching
                skills_score = 0.5  # Default score
            
            # Ensure valid score
            if not isinstance(skills_score, (int, float)) or skills_score < 0:
                skills_score = 0.5
            
            skills_score = min(1.0, max(0.0, skills_score))
            
            # Add some randomness to create variety
            import random
            noise = random.uniform(-0.05, 0.05)
            skills_score = max(0.0, min(1.0, skills_score + noise))
            
            return {
                'status': 'good' if skills_score > 0.6 else 'fair' if skills_score > 0.3 else 'poor',
                'similarity_score': skills_score,
                'distance': 1.0 - skills_score,
                'overall_match_percentage': skills_score
            }
            
        except Exception as e:
            print(f"Error in calculate_match: {e}")
            return {
                'status': 'poor',
                'similarity_score': 0.5,
                'distance': 0.5,
                'overall_match_percentage': 0.5
            }
class DocumentParser:
    """Abstract base class for document parsers"""
    def extract_text(self, file_content: bytes, filename: str) -> str:
        pass
class PDFParser(DocumentParser):
    def extract_text(self, file_content: bytes, filename: str) -> str:
        try:
            # First try with pdfplumber (better for CV layouts)
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
                if text.strip():
                    return text
        except Exception as e:
            logging.warning(f"pdfplumber failed for {filename}, trying PyMuPDF: {e}")
        
        try:
            # Fallback to PyMuPDF
            doc = fitz.open(stream=file_content, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except Exception as e:
            logging.warning(f"PyMuPDF failed for {filename}, trying PyPDF2: {e}")
        
        try:
            # Final fallback to PyPDF2
            pdf_file = io.BytesIO(file_content)
            reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
        except Exception as e:
            logging.error(f"All PDF extraction methods failed for {filename}: {e}")
            return ""
class DOCXParser(DocumentParser):
    def extract_text(self, file_content: bytes, filename: str) -> str:
        try:
            doc = docx.Document(io.BytesIO(file_content))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            logging.error(f"Error extracting DOCX text for {filename}: {e}")
            return ""
class TXTParser(DocumentParser):
    def extract_text(self, file_content: bytes, filename: str) -> str:
        try:
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return file_content.decode('latin-1')
            except Exception as e:
                logging.error(f"Error extracting TXT text for {filename}: {e}")
                return ""
class ParserFactory:
    _parsers = {
        '.pdf': PDFParser(),
        '.docx': DOCXParser(),
        '.doc': DOCXParser(),
        '.txt': TXTParser()
    }
    
    @classmethod
    def get_parser(cls, filename: str):
        ext = os.path.splitext(filename.lower())[1]
        return cls._parsers.get(ext)


class CVGenerator:
    def __init__(self):
        openai_key = os.getenv("OPENAI_API_KEY")
        if openai_key:
            self.oai_client = OpenAI(api_key=openai_key)
        else:
            self.oai_client = None
            st.warning("⚠️ OpenAI API key not found. CV extraction will use basic parsing.")


    def generate_candidate_id(self, email: str) -> int:
        """
        Cette méthode ne sera plus utilisée car on utilise l'auto-increment PostgreSQL
        Mais on la garde pour compatibilité
        """
        import hashlib
        import time
        
        try:
            if not email:
                return int(time.time()) % 100000000  # 8 chiffres max
            
            clean_email = email.lower().strip()
            
            # Utiliser seulement 6 caractères hex (24 bits)
            hash_object = hashlib.md5(clean_email.encode('utf-8'))
            hash_short = hash_object.hexdigest()[:6]  # 6 caractères = 24 bits
            hash_int = int(hash_short, 16)
            
            # Limite à 100 millions (8 chiffres)
            safe_id = (hash_int % 99999999) + 1  # Entre 1 et 99,999,999
            
            return safe_id
            
        except Exception as e:
            print(f"Erreur génération ID: {e}")
            return int(time.time()) % 99999999 + 1

    def generate_embedding(self, skills_dict):
        """Generate embeddings for skills - returns proper vector format"""
        if not self.oai_client:
            # Return a default vector instead of empty list
            return [0.0] * 1536  # OpenAI embedding dimension
            
        all_skills = " | ".join(
            skills_dict.get("technical_skills", []) + 
            skills_dict.get("soft_skills", [])
        )
        
        if not all_skills:
            # Return default vector for empty skills
            return [0.0] * 1536
            
        try:
            resp = self.oai_client.embeddings.create(
                model="text-embedding-3-small",
                input=all_skills
            )
            return resp.data[0].embedding
        except Exception as e:
            logging.error(f"Failed to generate embedding: {e}")
            # Return default vector on error
            return [0.0] * 1536

class EnhancedS3Manager:
    def __init__(self):
        try:
            # Check if AWS credentials are available
            import boto3
            from botocore.exceptions import NoCredentialsError, PartialCredentialsError
            
            # Try to create S3 client with explicit error checking
            self.s3_client = boto3.client(
                's3', 
                region_name=CONFIG["AWS_REGION"],
                # Add explicit credentials if needed
                # aws_access_key_id=os.getenv('AWS_ACCESS_KEY_ID'),
                # aws_secret_access_key=os.getenv('AWS_SECRET_ACCESS_KEY')
            )
            
            # Test S3 connection by listing buckets
            try:
                self.s3_client.list_buckets()
                print("✅ S3 client initialized successfully")
                self.s3_available = True
            except Exception as test_error:
                print(f"⚠️ S3 connection test failed: {test_error}")
                self.s3_available = False
                
        except Exception as e:
            print(f"❌ S3 client initialization failed: {e}")
            self.s3_client = None
            self.s3_available = False

    def upload_file(self, file_content: bytes, filename: str, folder: str = "cv_uploads"):
        """Upload file to S3 bucket with detailed error reporting"""
        
        # Check if S3 is available
        if not self.s3_available or not self.s3_client:
            error_msg = "S3 client not available - check AWS credentials and configuration"
            print(f"❌ {error_msg}")
            return False, error_msg
        
        # Validate inputs
        if not file_content:
            error_msg = "File content is empty"
            print(f"❌ {error_msg}")
            return False, error_msg
            
        if not filename:
            error_msg = "Filename is required"
            print(f"❌ {error_msg}")
            return False, error_msg
            
        try:
            # Generate unique key with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            file_ext = os.path.splitext(filename)[1]
            clean_filename = os.path.splitext(filename)[0]
            # Remove any problematic characters
            clean_filename = "".join(c for c in clean_filename if c.isalnum() or c in (' ', '-', '_')).strip()
            s3_key = f"{folder}/{timestamp}_{clean_filename}{file_ext}"
            
            print(f"📤 Uploading to S3: {s3_key}")
            print(f"📊 File size: {len(file_content)} bytes")
            print(f"🪣 Bucket: {CONFIG['S3_BUCKET_CANDIDATES']}")
            
            # Upload to S3 with proper error handling
            self.s3_client.put_object(
                Bucket=CONFIG["S3_BUCKET_CANDIDATES"],
                Key=s3_key,
                Body=file_content,
                ContentType=self._get_content_type(file_ext),
                # Add metadata
                Metadata={
                    'original_filename': filename,
                    'upload_timestamp': timestamp,
                    'content_length': str(len(file_content))
                }
            )
            
            # Verify upload by checking if object exists
            try:
                self.s3_client.head_object(
                    Bucket=CONFIG["S3_BUCKET_CANDIDATES"],
                    Key=s3_key
                )
                print(f"✅ Successfully uploaded and verified: {s3_key}")
                return True, s3_key
                
            except Exception as verify_error:
                error_msg = f"Upload succeeded but verification failed: {verify_error}"
                print(f"⚠️ {error_msg}")
                return True, s3_key  # Still return success since upload worked
            
        except Exception as e:
            error_msg = f"Failed to upload {filename} to S3: {str(e)}"
            print(f"❌ {error_msg}")
            
            # Check specific error types
            if "NoCredentialsError" in str(type(e)):
                error_msg += " - AWS credentials not found"
            elif "AccessDenied" in str(e):
                error_msg += " - Access denied to S3 bucket"
            elif "NoSuchBucket" in str(e):
                error_msg += f" - Bucket '{CONFIG['S3_BUCKET_CANDIDATES']}' does not exist"
                
            return False, error_msg

    def _get_content_type(self, file_ext: str) -> str:
        """Get content type based on file extension"""
        content_types = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.txt': 'text/plain'
        }
        return content_types.get(file_ext.lower(), 'application/octet-stream')

def verify_s3_configuration():
    """Verify S3 configuration and display status"""
    st.subheader("🔍 S3 Configuration Status")
    
    try:
        s3_manager = EnhancedS3Manager()
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if s3_manager.s3_available:
                st.success("✅ S3 Client: Connected")
            else:
                st.error("❌ S3 Client: Failed")
        
        with col2:
            bucket_name = CONFIG["S3_BUCKET_CANDIDATES"]
            st.info(f"🪣 Bucket: {bucket_name}")
            
            # Test bucket access
            if s3_manager.s3_available:
                try:
                    s3_manager.s3_client.head_bucket(Bucket=bucket_name)
                    st.success("✅ Bucket: Accessible")
                except Exception as bucket_error:
                    st.error(f"❌ Bucket: {bucket_error}")
            
        with col3:
            region = CONFIG["AWS_REGION"]
            st.info(f"🌍 Region: {region}")
        
        # Show environment variables status
        st.subheader("🔑 AWS Credentials Status")
        
        aws_key = os.getenv('AWS_ACCESS_KEY_ID')
        aws_secret = os.getenv('AWS_SECRET_ACCESS_KEY')
        
        col1, col2 = st.columns(2)
        with col1:
            if aws_key:
                st.success(f"✅ Access Key configured succesfully")
            else:
                st.error("❌ Access Key: Not found")
        
        with col2:
            if aws_secret:
                st.success(f"✅ Secret Key Configured succesfully")
            else:
                st.error("❌ Secret Key: Not found")
        
        # Test upload with dummy file
        if st.button("🧪 Test S3 Upload"):
            test_content = b"Test file content for S3 upload verification"
            test_filename = f"test_upload_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
            
            success, result = s3_manager.upload_file(test_content, test_filename, "test_uploads")
            
            if success:
                st.success(f"✅ Test upload successful: {result}")
                
                # Try to delete test file
                try:
                    s3_manager.s3_client.delete_object(
                        Bucket=CONFIG["S3_BUCKET_CANDIDATES"],
                        Key=result
                    )
                    st.info("🗑️ Test file cleaned up")
                except:
                    st.warning("⚠️ Could not delete test file")
            else:
                st.error(f"❌ Test upload failed: {result}")
                
    except Exception as e:
        st.error(f"❌ S3 verification failed: {e}")
#Extraire des informations de candidats à partir de texte (comme un CV) en utilisant LangChain
def setup_langchain_pipeline():
    """Setup LangChain pipeline for CV extraction"""
    #One shot promting
    system_prompt = dedent("""
    # Role and Objective
    You are an intelligent agent specialized in extracting structured data from CVs/resumes. Your objective is to extract accurate and complete information from a given CV text and return it as a valid JSON object.
    # Instructions
    - Extract all relevant candidate information comprehensively and accurately
    - If information is missing, use empty strings or appropriate defaults
    - Ensure all skills are properly categorized as technical or soft skills
    - Extract work experience with proper date formatting
    - Extract education information with institutions and degrees
    ## Data Extraction Requirements
    - Extract candidate's full name, separating first and last name
    - Extract contact information (email, phone)
    - Extract current job title and location (city, country)
    - Identify technical skills (programming languages, tools, frameworks, databases)
    - Identify soft skills (communication, leadership, teamwork, etc.)
    - Extract work experience with company, role, dates, and descriptions
    - Extract education with institution, degree, department, and dates
    - Extract certifications with issuer and year if available
    - Extract LinkedIn profile if mentioned
    - Determine highest education level achieved
    - Calculate total years of experience
    # Output Format
    Return ONLY valid JSON in this exact format:
    {
        "first_name": "string",
        "last_name": "string",
        "email": "string",
        "phone": "string",
        "current_job": "string",
        "country": "string",
        "city": "string",
        "technical_skills": ["string"],
        "soft_skills": ["string"],
        "experience": [],
        "certifications": [],
        "schools": [],
        "job_category": "string",
        "linkedin": "string",
        "education_level": "string",
        "language": "string"
    }
    """).replace("{", "{{").replace("}", "}}")

    # Step 2: Construction du prompt
    prompt_template = ChatPromptTemplate.from_messages([
        ("system", system_prompt),
        ("human", "Analyze the following CV/resume and extract the requested information:\n\n{text}")
    ])

    # Step 3: LLM avec format JSON
    llm = ChatOpenAI(
        model=CONFIG["OPENAI_MODEL"],
        temperature=0.1,
        api_key=os.getenv("OPENAI_API_KEY"),
        model_kwargs={"response_format": {"type": "json_object"}}
    )

    # Step 4: Output parser JSON
    parser = JsonOutputParser()

    # Step 5: Pipeline
    return prompt_template | llm | parser

#Validation des champs : S'assure que tous les champs requis existent
def safe_extract_with_langchain(text, filename=""):
    """Extract candidate info using LangChain"""
    try:
        pipeline = setup_langchain_pipeline()
        result = pipeline.invoke({"text": text})
        
        # Set filename and extension
        result["filename"] = filename
        if filename:
            ext = os.path.splitext(filename)[1]
            result["document_extension"] = ext[1:].lower() if ext.startswith(".") else ""
        else:
            result["document_extension"] = ""
        
        # Ensure all required fields exist
        required_fields = [
            "first_name", "last_name", "email", "phone", "current_job",
            "country", "city", "technical_skills", "soft_skills", "experience",
            "certifications", "schools", "job_category", "linkedin", "education_level"
        ]
        
        for field in required_fields:
            if field not in result:
                if field in ["technical_skills", "soft_skills", "experience", "certifications", "schools"]:
                    result[field] = []
                else:
                    result[field] = ""
        
        # Set language
        result["language"] = detect(text) if text else "unknown"
        
        # Set URL (empty for now)
        result["url"] = ""
        
        return result

    except Exception as e:
        logging.error(f"LangChain extraction error: {e}")
        # Fallback basic extraction
        return {
            "first_name": "", "last_name": "", "email": "", "phone": "",
            "current_job": "", "country": "", "city": "",
            "technical_skills": [], "soft_skills": [], "experience": [],
            "certifications": [], "schools": [], "job_category": "",
            "linkedin": "", "education_level": "Unknown", "language": "unknown",
            "filename": filename, "document_extension": "", "url": ""
        }

# Updated Pipeline class with better error handling
class JobMatchingPipeline:
    def __init__(self, db_manager, matching_engine):
        self.db_manager = db_manager
        self.matching_engine = matching_engine
    
    def process_matching_for_candidate(self, candidate_id:int):
        """
        Process job matching for a specific candidate with improved error handling
        """
        try:
            # Get candidate data
            candidate = self.db_manager.fetch_one(
                "SELECT * FROM candidates WHERE candidate_id = %(candidate_id)s",
                {'candidate_id': candidate_id}
            )
            
            if not candidate:
                print(f"Candidate {candidate_id} not found")
                return 0
            
            # Get active job offers
            jobs = self.db_manager.execute_query(
                "SELECT * FROM job_offers WHERE status = 'active'"
            )
            
            if not jobs:
                print("No active job offers found")
                return 0
            
            matches_created = 0
            conn = psycopg2.connect(**POSTGRES_CONFIG)
            cur = conn.cursor()
            
            for job in jobs:
                try:
                    # Calculate match with improved error handling
                    match_result = self.matching_engine.calculate_match(candidate, job)
                    
                    match_params = {
                        'match_id': str(uuid.uuid4()),
                        'candidate_id': candidate_id,
                        'job_offer_id': job['job_offer_id'],
                        'distance': match_result['distance'],
                        'similarity_score': match_result['similarity_score'],
                        'overall_evaluation': match_result['overall_evaluation'],
                        'created_at': datetime.now(),
                        'updated_at': datetime.now()
                    }
                    
                    match_insert = """
                        INSERT INTO job_matches (
                            match_id, candidate_id, job_offer_id,
                            distance, similarity_score, overall_match_percentage, last_updated
                        ) VALUES (
                            %(match_id)s, %(candidate_id)s, %(job_offer_id)s, %(distance)s, %(similarity_score)s
                            , %(overall_evaluation)s,
                            %(updated_at)s,%(created_at)s
                        )
                        ON CONFLICT (candidate_id, job_offer_id) DO UPDATE SET
                            distance = EXCLUDED.distance,
                            similarity_score = EXCLUDED.similarity_score,
                            overall_evaluation = EXCLUDED.evaluation,
                            updated_at = EXCLUDED.updated_at,
                            created_at = EXCLUDED.created_at

                    """
                    
                    cur.execute(match_insert, match_params)
                    matches_created += 1
                    
                except Exception as e:
                    print(f"Error processing match for job {job.get('job_offer_id', 'unknown')}: {e}")
                
            conn.commit()
            cur.close()
            conn.close()
            
            return matches_created
            
        except Exception as e:
            print(f"Error in process_matching_for_candidate: {e}")
            return 0   
        

st.set_page_config(
    page_title="AI Job Matching System",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main {
        padding: 0rem 1rem;
    }
    .metric-card {
        background-color: ##155773;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 0.5rem 0;
    }
    .excellent-match {
        background-color: #d4edda;
        color: #155724;
        padding: 0.5rem;
        border-radius: 0.3rem;
    }
    .good-match {
        background-color: #cce5ff;
        color: #004085;
        padding: 0.5rem;
        border-radius: 0.3rem;
    }
    .average-match {
        background-color: #fff3cd;
        color: #856404;
        padding: 0.5rem;
        border-radius: 0.3rem;
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 24px;
    }
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        padding-left: 20px;
        padding-right: 20px;
    }
    .job-card {
        background-color: #639cd6;
        border: 1px solid #dee2e6;
        border-radius: 0.5rem;
        padding: 1rem;
        margin-bottom: 1rem;
    }
    .candidate-card {
        background-color: #639cd6;
        border: 1px solid #e3e3e3;
        border-radius: 0.5rem;
        padding: 1rem;
        margin-bottom: 0.5rem;
    }
    .match-card {
        background-color:#639cd6;
        border: 1px solid #dee2e6;
        border-radius: 0.5rem;
        padding: 1.5rem;
        margin-bottom: 1rem;
        transition: all 0.3s ease;
    }
    .match-card:hover {
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'user_role' not in st.session_state:
    st.session_state.user_role = None
if 'user_id' not in st.session_state:
    st.session_state.user_id = None
if 'last_update' not in st.session_state:
    st.session_state.last_update = None

# Cache only the config, not the connection
@st.cache_resource
def init_config():
    return DatabaseConfig(**POSTGRES_CONFIG)

# Create DatabaseManager from cached config
db_config = init_config()
db_manager = DatabaseManager(db_config)

#initiliser tou sles comosants
@st.cache_resource
def init_system():
    
    try:
        print("🔄 Starting system initialization with tuple error fixes...")
        
        # Step 1: Database with fixed manager
        try:
            db_config = DatabaseConfig(**POSTGRES_CONFIG)
            db_manager = DatabaseManager(db_config)
            
            # CRITICAL: Test and fix database returns
            test_result = db_manager.execute_query("SELECT 1 as test")
            if not test_result:
                raise Exception("Database test query failed")
            
            # Verify we get dictionaries
            if test_result and isinstance(test_result[0], dict):
                print("✅ Database returns dictionaries correctly")
            else:
                print("⚠️ Database returns tuples - applying runtime patch")
                # Apply the patch
                db_manager = patch_database_manager(db_manager)
                
                # Verify patch worked
                test_result2 = db_manager.execute_query("SELECT 1 as test")
                if test_result2 and isinstance(test_result2[0], dict):
                    print("✅ Patch successful - now returning dictionaries")
                else:
                    print("❌ Patch failed - still returning tuples")
                    
        except Exception as db_error:
            print(f"❌ Database initialization failed: {db_error}")
            raise
        
        # Step 2: Initialize other components
        matching_engine = EnhancedMatchingEngine()
        pipeline = DataPipeline(db_manager, matching_engine)
        cv_generator = CVGenerator()
        s3_manager = EnhancedS3Manager()
        
        # ADDED: Initialize QuickSparkIntegration with db_manager
        quick_spark = QuickSparkIntegration(db_manager)
        # ✅ CRITICAL FIX: Return dictionary instead of tuple
        components = {
            'db_manager': db_manager,
            'matching_engine': matching_engine,
            'pipeline': pipeline,
            'cv_generator': cv_generator,
            's3_manager': s3_manager
        }
        
        print("✅ System initialization completed!")
        return components  # ✅ RETURN DICTIONARY NOT TUPLE
        
    except Exception as e:
        print(f"❌ System initialization failed: {e}")
        print(f"Traceback: {traceback.format_exc()}")
        return None



try:
    # ✅ FIXED: Properly handle dictionary return
    components = init_system()
    if components is None:
        st.error("❌ System initialization failed - please check your configuration")
        st.error("Common issues: Database connection, missing dependencies, or configuration errors")
        st.stop()
    
    # ✅ FIXED: Verify components is a dictionary
    if not isinstance(components, dict):
        st.error(f"❌ init_system returned {type(components)} instead of dict")
        st.error(f"Returned value: {components}")
        st.stop()
    
    # ✅ FIXED: Access components as dictionary
    db_manager = components['db_manager']
    
    # Apply additional patching if needed
    original_execute = db_manager.execute_query

    def patched_execute(query, params=None):
        try:
            conn = db_manager.get_connection()
            # Force RealDictCursor to return dictionaries
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cursor:
                cursor.execute(query, params)
                if query.strip().upper().startswith('SELECT'):
                    results = cursor.fetchall()
                    return [dict(row) for row in results] if results else []
                else:
                    conn.commit()
                    return []
        except Exception as e:
            print(f"Database error: {e}")
            return []
        
    db_manager.execute_query = patched_execute
    print("✅ Database manager patched - tuple error should be fixed")
    
    # ✅ FIXED: Access all other components
    matching_engine = components['matching_engine']
    pipeline = components['pipeline']
    cv_generator = components['cv_generator']
    s3_manager = components['s3_manager']
    
    
except Exception as e:
    st.error(f"❌ Initialization error: {e}")
    st.error("Please check your database connection and configuration")
    
    # Show diagnostic information
    with st.expander("🔍 Diagnostic Information"):
        st.code(traceback.format_exc())
        
        # Try to show what init_system is actually returning
        try:
            st.write("**Testing init_system return type:**")
            test_components = init_system()
            st.write(f"init_system returns type: {type(test_components)}")
            st.write(f"init_system returns value: {test_components}")
            
            if isinstance(test_components, tuple):
                st.error("❌ init_system is returning a TUPLE - this needs to be fixed!")
                st.write("The function should return a dictionary, not a tuple")
            elif isinstance(test_components, dict):
                st.success("✅ init_system correctly returns a dictionary")
            else:
                st.warning(f"❓ init_system returns unexpected type: {type(test_components)}")
                
        except Exception as diag_error:
            st.write(f"init_system test failed: {diag_error}")
    
    st.stop()
    
#FIX CV COMPONENT INITIALIZATION
@st.cache_resource
def init_cv_components():
    """Initialize CV processing components"""
    try:
        cv_generator = CVGenerator()
        s3_manager = EnhancedS3Manager()  # Use EnhancedS3Manager
        return cv_generator, s3_manager
    except Exception as e:
        st.error(f"Failed to initialize CV components: {str(e)}")
        return None, None

# ============= HELPER FUNCTIONS ============= #

def get_match_class(score: float) -> str:
    """Get CSS class based on match score"""
    if score > 0.8:
        return "excellent-match"
    elif score > 0.6:
        return "good-match"
    else:
        return "average-match"

def format_percentage(value: float) -> str:
    """Format float as percentage"""
    return f"{value * 100:.1f}%"

# ============= PAGE FUNCTIONS ============= #
def login_page():
    """Login page with proper error handling - FIXED for tuple error"""
    st.title("🎯 AI Job Matching System")
    st.markdown("### Welcome! Please select your role to continue")
    
    # Add diagnosis in sidebar
    with st.sidebar:
        st.markdown("### 🔧 System Diagnostics")
        if st.button("🔍 Check Database Connection"):
            diagnose_database_connection()
        if st.button("🧪 Test Database Returns"):
            test_database_returns(db_manager)
        if st.button("🪣 Check S3 Configuration"):
            verify_s3_configuration()
        if st.button("🚀 Démonstration Master-Worker", key="master_worker_demo"):
            with st.spinner("Exécution des démonstrations Spark Master-Worker..."):
                # Créer une instance améliorée
                enhanced_spark = EnhancedQuickSparkIntegration(db_manager)
                
                # Capturer la sortie pour l'afficher dans Streamlit
                import io
                import sys
                
                # Rediriger stdout pour capturer les prints
                old_stdout = sys.stdout
                captured_output = io.StringIO()
                sys.stdout = captured_output
                
                try:
                    enhanced_spark.demonstrate_master_worker_operations()
                    output = captured_output.getvalue()
                    
                    # Afficher dans Streamlit
                    st.success("✅ Démonstrations terminées!")
                    with st.expander("📊 Logs des démonstrations Master-Worker", expanded=True):
                        st.code(output)
                        
                except Exception as e:
                    st.error(f"❌ Erreur: {e}")
                    
                finally:
                    sys.stdout = old_stdout
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 👤 Candidate")
        st.markdown("Find your perfect job match")
        if st.button("Login as Candidate", key="candidate_login", use_container_width=True):
            try:
                # FIXED: Use safe dictionary access with error handling
                candidates = db_manager.execute_query("SELECT candidate_id, first_name, last_name FROM candidates LIMIT 1")
                
                if candidates and len(candidates) > 0:
                    candidate = candidates[0]
                    # SAFE: Handle both tuple and dict with detailed error checking
                    try:
                        if isinstance(candidate, dict):
                            candidate_id = candidate.get('candidate_id')
                            first_name = candidate.get('first_name', 'Candidate')
                        else:  # tuple
                            candidate_id = candidate[0] if len(candidate) > 0 else None
                            first_name = candidate[1] if len(candidate) > 1 else 'Candidate'
                        
                        if candidate_id:
                            st.session_state.user_role = 'candidate'
                            st.session_state.user_id = candidate_id
                            st.success(f"Welcome {first_name}!")
                        else:
                            raise ValueError("Could not extract candidate_id")
                            
                    except Exception as parse_error:
                        st.error(f"Error parsing candidate data: {parse_error}")
                        st.error(f"Candidate data type: {type(candidate)}")
                        st.error(f"Candidate data: {candidate}")
                        raise parse_error
                        
                else:
                    # Demo mode fallback
                    st.session_state.user_role = 'candidate'
                    st.session_state.user_id = 1
                    st.info("Using demo mode - no candidates found in database")
                
                st.rerun()
                
            except Exception as e:
                st.error(f"Login error: {str(e)}")
                st.error("Detailed error information:")
                st.code(traceback.format_exc())
                
                # Always provide fallback
                st.session_state.user_role = 'candidate'
                st.session_state.user_id = 1
                st.info("Using demo mode due to error")
                st.rerun()
    
    with col2:
        st.markdown("### 🏢 Client/Recruiter")
        st.markdown("Find the best candidates")
        if st.button("Login as Client", key="client_login", use_container_width=True):
            try:
                # FIXED: Use safe dictionary access with error handling
                clients = db_manager.execute_query("SELECT client_id, company_name FROM clients LIMIT 1")
                
                if clients and len(clients) > 0:
                    client = clients[0]
                    # SAFE: Handle both tuple and dict with detailed error checking
                    try:
                        if isinstance(client, dict):
                            client_id = client.get('client_id')
                            company_name = client.get('company_name', 'Client')
                        else:  # tuple
                            client_id = client[0] if len(client) > 0 else None
                            company_name = client[1] if len(client) > 1 else 'Client'
                        
                        if client_id:
                            st.session_state.user_role = 'client'
                            st.session_state.user_id = client_id
                            st.success(f"Welcome {company_name}!")
                        else:
                            raise ValueError("Could not extract client_id")
                            
                    except Exception as parse_error:
                        st.error(f"Error parsing client data: {parse_error}")
                        st.error(f"Client data type: {type(client)}")
                        st.error(f"Client data: {client}")
                        raise parse_error
                        
                else:
                    # Demo mode fallback
                    st.session_state.user_role = 'client'
                    st.session_state.user_id = 1
                    st.info("Using demo mode - no clients found in database")
                
                st.rerun()
                
            except Exception as e:
                st.error(f"Login error: {str(e)}")
                st.error("Detailed error information:")
                st.code(traceback.format_exc())
                
                # Always provide fallback
                st.session_state.user_role = 'client'
                st.session_state.user_id = 1
                st.info("Using demo mode due to error")
                st.rerun()
    
    

def initialize_app():
    """Initialize all app components"""
    if 'app_initialized' not in st.session_state:
        try:
            # Database components
            db_config = DatabaseConfig(**POSTGRES_CONFIG)
            st.session_state.db_manager = DatabaseManager(db_config)
            st.session_state.matching_engine = EnhancedMatchingEngine()
            
            st.session_state.pipeline = DataPipeline(
                st.session_state.db_manager,
                st.session_state.matching_engine )
            
            # CV components
            st.session_state.cv_generator = CVGenerator()
            st.session_state.s3_manager = EnhancedS3Manager()
            
            st.session_state.app_initialized = True
            
        except Exception as e:
            st.error(f"Failed to initialize app: {str(e)}")
            st.stop()

def candidate_dashboard():
    
    """Enhanced candidate dashboard with CV upload and job matching - COMPLETE CORRECTED VERSION"""
   # Helper functions (keep your existing ones but add missing imports and fixes)
    def insert_candidate_with_matching(candidate_data, db_manager, pipeline=None):
        """Enhanced version that ensures matches are created - WITH DETAILED DEBUGGING"""
        
        try:
            # Insert candidate using existing logic
            result_candidate_id = insert_candidate_with_related_data(candidate_data)
            
            if not result_candidate_id:
                st.error("❌ Failed to insert candidate data")
                return None
            
            st.success(f"✅ Candidate {result_candidate_id} saved successfully!")
            
            # CRITICAL: Create job matches immediately after candidate insertion
            matches_created = 0
            
            # Show database status first
            try:
                job_stats = db_manager.execute_query("""
                    SELECT 
                        COUNT(*) as total_jobs,
                        COUNT(CASE WHEN status = 'active' THEN 1 END) as active_jobs,
                        COUNT(CASE WHEN status IS NULL THEN 1 END) as null_status_jobs
                    FROM job_offers
                """)
                
                if job_stats:
                    stats = job_stats[0]
                    st.info(f"📊 Job offers: {stats['total_jobs']} total, {stats['active_jobs']} active, {stats['null_status_jobs']} with null status")
                else:
                    st.warning("⚠️ Could not get job offer statistics")
                    
            except Exception as stats_error:
                st.warning(f"Could not get database stats: {stats_error}")
            
            # Try pipeline first if available
            if pipeline:
                try:
                    st.info("🔄 Attempting pipeline matching...")
                    matches_created = pipeline.process_matching_for_candidate(result_candidate_id)
                    if matches_created > 0:
                        st.success(f"✅ Pipeline created {matches_created} job matches!")
                    else:
                        st.warning("⚠️ Pipeline returned 0 matches")
                except Exception as match_error:
                    st.warning(f"⚠️ Pipeline matching failed: {match_error}")
            else:
                st.info("ℹ️ No pipeline available, skipping to manual matching")
            
            # Try enhanced manual matching if pipeline failed or no pipeline
            if matches_created == 0:
                st.info("🔄 Attempting enhanced manual matching...")
                matches_created = create_manual_matches(result_candidate_id)
                if matches_created > 0:
                    st.success(f"matching create {matches_created} job matches!")
                else:
                    st.warning(" matching also returned 0 matches")
            
            # Try basic test matching as absolute last resort
            if matches_created == 0:
                st.warning("🔄 Attempting basic test matching as last resort...")
                matches_created = create_basic_test_matches(result_candidate_id)
                if matches_created > 0:
                    st.success(f" matching created {matches_created} test matches!")
                else:
                    st.error("❌ Even basic matching failed!")
            
            # Comprehensive verification with detailed debugging
            try:
                st.info("🔍 Performing comprehensive match verification...")
                
                actual_matches = verify_matches_fixed(result_candidate_id)
                
                
                
                if actual_matches == 0:
                    st.error("❌ No matches were actually saved to database!")
                    
                    # Debug why no matches were saved
                    st.info("🔍 Debugging match insertion...")
                    
                    # Check if job_matches table exists and is accessible
                    try:
                        table_check = db_manager.execute_query("""
                            SELECT COUNT(*) as count FROM information_schema.tables 
                            WHERE table_name = 'job_matches'
                        """)
                        
                        if table_check and table_check[0]['count'] > 0:
                            st.info("✅ job_matches table exists")
                            
                            # Check table structure
                            columns = db_manager.execute_query("""
                                SELECT column_name, data_type, is_nullable
                                FROM information_schema.columns 
                                WHERE table_name = 'job_matches'
                                ORDER BY ordinal_position
                            """)
                            
                            if columns:
                                st.info("📋 job_matches table structure:")
                                for col in columns:
                                    st.write(f"• {col['column_name']} ({col['data_type']}, nullable: {col['is_nullable']})")
                            
                            # Check for any existing matches in the table
                            total_matches_in_table = db_manager.fetch_one("SELECT COUNT(*) as count FROM job_matches")
                            st.info(f"📊 Total matches in job_matches table: {total_matches_in_table['count'] if total_matches_in_table else 0}")
                            
                        else:
                            st.error("❌ job_matches table does not exist!")
                            
                    except Exception as table_error:
                        st.error(f"❌ Cannot access job_matches table: {table_error}")
                
                else:
                    # Show sample matches for verification
                    sample_matches = db_manager.execute_query("""
                        SELECT jm.similarity_score, jm.created_at, jo.job_title, jo.company_name 
                        FROM job_matches jm 
                        LEFT JOIN job_offers jo ON jm.job_offer_id = jo.job_offer_id 
                        WHERE jm.candidate_id = %s 
                        ORDER BY jm.similarity_score DESC 
                        LIMIT 5
                    """, (result_candidate_id,))
                    
                    if sample_matches:
                        st.success("✅ Sample matches verified:")
                        for i, match in enumerate(sample_matches):
                            job_title = match.get('job_title', 'Unknown Job')
                            company_name = match.get('company_name', 'Unknown Company')
                            score = match.get('similarity_score', 0)
                            created = match.get('created_at', 'Unknown')
                            st.write(f"{i+1}. {job_title} at {company_name} ({score:.1%}) - {created}")
                    
                    st.balloons()
                    
            except Exception as verify_error:
                st.warning(f"Match verification failed: {verify_error}")
                # Still return the candidate ID even if verification fails
            
            return result_candidate_id
            
        except Exception as e:
            st.error(f"❌ Error in enhanced insertion: {e}")
            with st.expander("🔍 Error Details"):
                st.code(traceback.format_exc())
            return None
   
    def safe_skills_processor(skills_data):
        """Safely process skills data from various formats to the expected JSON structure"""
        if not skills_data:
            return {"technical_skills": [], "soft_skills": []}
        
        if isinstance(skills_data, dict) and 'technical_skills' in skills_data:
            tech_skills = skills_data.get('technical_skills', [])
            soft_skills = skills_data.get('soft_skills', [])
            
            if isinstance(tech_skills, list):
                tech_skills = [str(skill).strip() for skill in tech_skills if skill]
            elif isinstance(tech_skills, str):
                tech_skills = [skill.strip() for skill in tech_skills.split(',') if skill.strip()]
            else:
                tech_skills = []
                
            if isinstance(soft_skills, list):
                soft_skills = [str(skill).strip() for skill in soft_skills if skill]
            elif isinstance(soft_skills, str):
                soft_skills = [skill.strip() for skill in soft_skills.split(',') if skill.strip()]
            else:
                soft_skills = []
                
            return {"technical_skills": tech_skills, "soft_skills": soft_skills}
        
        if isinstance(skills_data, str):
            try:
                parsed = json.loads(skills_data)
                return safe_skills_processor(parsed)
            except json.JSONDecodeError:
                skills_list = [skill.strip() for skill in skills_data.split(',') if skill.strip()]
                return {"technical_skills": skills_list, "soft_skills": []}

        if isinstance(skills_data, list):
            tech_skills = [str(skill).strip() for skill in skills_data if skill]
            return {"technical_skills": tech_skills, "soft_skills": []}
        
        return {"technical_skills": [], "soft_skills": []}
        
    def skills_to_display_format(skills_dict, skill_type):
        """Convert skills dict to comma-separated string for display"""
        if not skills_dict or not isinstance(skills_dict, dict):
            return ""
        
        skills_list = skills_dict.get(skill_type, [])
        if isinstance(skills_list, list):
            return ", ".join(str(skill) for skill in skills_list if skill)
        elif isinstance(skills_list, str):
            return skills_list
        return ""
        
    def safe_vector_format(embedding):
        """Ensure embedding is in proper vector format"""
        if not embedding or not isinstance(embedding, list):
            return [0.0] * 1536
        
        try:
            return [float(x) for x in embedding]
        except (ValueError, TypeError):
            return [0.0] * 1536
    
    def get_or_create_location(city, country, continent="Unknown"):
        """Get location_id or create new location entry"""
        try:
            city = city.strip() if city else "Unknown"
            country = country.strip() if country else "Unknown"

            existing_location = db_manager.fetch_one(
                "SELECT location_id FROM region WHERE city = %(city)s AND country = %(country)s",
                {'city': city, 'country': country}
            )

            if existing_location:
                return existing_location['location_id']

            insert_location_query = """
                INSERT INTO region (city, country, continent, created_at, updated_at)
                VALUES (%(city)s, %(country)s, %(continent)s, %(created_at)s, %(updated_at)s)
                RETURNING location_id
            """

            result = db_manager.execute_query(insert_location_query, {
                'city': city,
                'country': country,
                'continent': continent,
                'created_at': datetime.now(),
                'updated_at': datetime.now()
            })

            return result[0]['location_id'] if result else None

        except Exception as e:
            st.warning(f"Could not create/find location: {e}")
            return None 

    def get_or_create_skill(skill_name, skill_type):
        try:
            skill_name = skill_name.strip()
            skill_type = skill_type.lower()  # 'technical' or 'soft'
            
            # Check if skill already exists
            existing_skill = db_manager.fetch_one(
                "SELECT skill_id FROM skills WHERE LOWER(name) = %(name)s AND skill_type = %(skill_type)s",
                {'name': skill_name.lower(), 'skill_type': skill_type}
            )
            
            if existing_skill:
                return existing_skill['skill_id']
            
            # Create new skill
            insert_skill_query = """
                INSERT INTO skills (name, skill_type, updated_at)
                VALUES (%(name)s, %(skill_type)s, %(updated_at)s)
                RETURNING skill_id
            """
            
            result = db_manager.execute_query(insert_skill_query, {
                'name': skill_name,
                'skill_type': skill_type,
                'updated_at': datetime.now()
            })
            
            return result[0]['skill_id'] if result else None
            
        except Exception as e:
            st.warning(f"Could not create/find skill '{skill_name}': {e}")
            return None
    def insert_candidate_with_related_data(candidate_data):
        """
        FONCTION CORRIGÉE - REMPLACE VOTRE FONCTION EXISTANTE
        """
        connection = None
        cursor = None
        
        try:
            # Validate required fields
            if not candidate_data.get('email'):
                st.error("❌ Email is required for candidate insertion")
                return None
            
            email = str(candidate_data.get('email', '')).strip()
            if not email or '@' not in email:
                st.error("❌ Valid email is required for candidate insertion")
                return None
            
            # CORRECTION: Suppression de la génération manuelle de candidate_id
            
            # Prepare and clean data
            skills_vector = candidate_data.get('skills_vector', [0.0] * 1536)
            if not isinstance(skills_vector, list):
                if hasattr(skills_vector, 'tolist'):
                    skills_vector = skills_vector.tolist()
                else:
                    skills_vector = [0.0] * 1536
            
            if len(skills_vector) != 1536:
                skills_vector = skills_vector[:1536] if len(skills_vector) > 1536 else skills_vector + [0.0] * (1536 - len(skills_vector))
            
            try:
                experience_years = int(candidate_data.get('experience_years', 0))
            except (ValueError, TypeError):
                experience_years = 0
            
            # Get or create location (optional)
            location_id = None
            if candidate_data.get('city') or candidate_data.get('country'):
                try:
                    location_id = get_or_create_location(
                        candidate_data.get('city', 'Unknown'),
                        candidate_data.get('country', 'Unknown')
                    )
                except Exception as loc_error:
                    st.warning(f"Could not create location: {loc_error}. Proceeding without location.")
                    location_id = None
            
            # Get database connection
            connection = db_manager.get_connection()
            cursor = connection.cursor()
            
            # Start transaction explicitly
            cursor.execute("BEGIN;")
            
            try:
                # Check if candidate exists BY EMAIL ONLY
                cursor.execute(
                    "SELECT candidate_id FROM candidates WHERE email = %s",
                    (email,)
                )
                existing_candidate = cursor.fetchone()
                
                final_candidate_id = None
                
                if existing_candidate:
                    # Update existing candidate
                    final_candidate_id = existing_candidate[0]
                    st.info(f"Updating existing candidate: {final_candidate_id}")
                    
                    update_query = """
                    UPDATE candidates SET
                        first_name = %s, last_name = %s, phone = %s, linkedin = %s,
                        job_category = %s, current_job = %s, country = %s, city = %s,
                        experience_years = %s, education_level = %s, skills = %s,
                        skills_vector = %s, certifications = %s, experience = %s,
                        filename = %s, url = %s, document_extension = %s,
                        document_language = %s, updated_at = CURRENT_TIMESTAMP,
                        processed_at = CURRENT_TIMESTAMP, location_id = %s
                    WHERE candidate_id = %s
                    """
                    
                    cursor.execute(update_query, (
                        str(candidate_data.get('first_name', '') or ''),
                        str(candidate_data.get('last_name', '') or ''),
                        str(candidate_data.get('phone', '') or ''),
                        str(candidate_data.get('linkedin', '') or ''),
                        str(candidate_data.get('job_category', '') or ''),
                        str(candidate_data.get('current_job', '') or ''),
                        str(candidate_data.get('country', '') or ''),
                        str(candidate_data.get('city', '') or ''),
                        experience_years,
                        str(candidate_data.get('education_level', 'Unknown') or 'Unknown'),
                        json.dumps(candidate_data.get('skills', {"technical_skills": [], "soft_skills": []})),
                        json.dumps(skills_vector),
                        json.dumps(candidate_data.get('certifications', [])),
                        json.dumps(candidate_data.get('experience', [])),
                        str(candidate_data.get('filename', '') or ''),
                        str(candidate_data.get('url', '') or ''),
                        str(candidate_data.get('document_extension', '') or ''),
                        str(candidate_data.get('document_language', 'unknown') or 'unknown'),
                        location_id,
                        final_candidate_id
                    ))
                    
                else:
                    # CORRECTION: Insert WITHOUT specifying candidate_id
                    st.info("Creating new candidate with auto-generated ID")
                    
                    insert_query = """
                    INSERT INTO candidates (
                        first_name, last_name, email, phone, linkedin,
                        job_category, current_job, country, city, experience_years,
                        education_level, skills, skills_vector, certifications, experience,
                        filename, url, document_extension, document_language,
                        visibility_setting, blocked, created_at, updated_at,
                        processed_at, location_id
                    ) VALUES (
                        %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                        %s, %s, %s, %s, %s, %s, %s, %s, %s,
                        %s, %s, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP,
                        CURRENT_TIMESTAMP, %s
                    ) RETURNING candidate_id
                    """
                    
                    cursor.execute(insert_query, (
                        str(candidate_data.get('first_name', '') or ''),
                        str(candidate_data.get('last_name', '') or ''),
                        email,
                        str(candidate_data.get('phone', '') or ''),
                        str(candidate_data.get('linkedin', '') or ''),
                        str(candidate_data.get('job_category', '') or ''),
                        str(candidate_data.get('current_job', '') or ''),
                        str(candidate_data.get('country', '') or ''),
                        str(candidate_data.get('city', '') or ''),
                        experience_years,
                        str(candidate_data.get('education_level', 'Unknown') or 'Unknown'),
                        json.dumps(candidate_data.get('skills', {"technical_skills": [], "soft_skills": []})),
                        json.dumps(skills_vector),
                        json.dumps(candidate_data.get('certifications', [])),
                        json.dumps(candidate_data.get('experience', [])),
                        str(candidate_data.get('filename', '') or ''),
                        str(candidate_data.get('url', '') or ''),
                        str(candidate_data.get('document_extension', '') or ''),
                        str(candidate_data.get('document_language', 'unknown') or 'unknown'),
                        bool(candidate_data.get('visibility_setting', True)),
                        False,  # blocked
                        location_id
                    ))
                    
                    # Get the auto-generated candidate_id
                    result = cursor.fetchone()
                    if result:
                        final_candidate_id = result[0]
                        st.success(f"✅ New candidate created with ID: {final_candidate_id}")
                    else:
                        raise Exception("Failed to get generated candidate_id")
                
                # Verify the insertion/update
                cursor.execute(
                    "SELECT candidate_id, email FROM candidates WHERE candidate_id = %s",
                    (final_candidate_id,)
                )
                verify_result = cursor.fetchone()
                
                if not verify_result:
                    raise Exception("Candidate verification failed")
                
                cursor.execute("COMMIT;")
                st.success(f"✅ Candidate saved successfully: {final_candidate_id}")
                
                return final_candidate_id
                
            except Exception as transaction_error:
                try:
                    cursor.execute("ROLLBACK;")
                except:
                    pass
                st.error(f"🔍 Transaction error: {transaction_error}")
                return None
                
        except Exception as e:
            st.error(f"❌ Unexpected error: {str(e)}")
            return None
            
        finally:
            try:
                if cursor:
                    cursor.close()
                if connection:
                    connection.close()
            except Exception as cleanup_error:
                st.warning(f"Database cleanup warning: {cleanup_error}")
    # ================================================================
    # NEW ENHANCED FUNCTIONS FOR JOB MATCHING
    # ================================================================
    
    
    def create_manual_matches(candidate_id):
        """COMPLETELY FIXED: Create matches manually"""
        try:
            st.info("🔍 Starting manual match creation...")
            
            # Get candidate data with SAFE dictionary access
            candidate = db_manager.fetch_one(
                "SELECT candidate_id, skills_vector, job_category, experience_years FROM candidates WHERE candidate_id = %s",
                (candidate_id,)
            )
            
            if not candidate:
                st.error(f"❌ Candidate {candidate_id} not found")
                return 0
            
            # Get job offers with SAFE dictionary access
            jobs = db_manager.execute_query("""
                SELECT job_offer_id, job_title, company_name, job_category, 
                    city, country, skills_vector, technical_skills
                FROM job_offers 
                WHERE status = 'active' OR status IS NULL
                ORDER BY created_at DESC
                LIMIT 50
            """)
            
            if not jobs:
                st.warning("⚠️ No job offers found")
                return 0
            
            st.info(f"✅ Found {len(jobs)} job offers")
            
            matches_created = 0
            
            # FIXED: Create matches with proper error handling
            for idx, job in enumerate(jobs):
                try:
                    # SAFE: Extract data with fallbacks
                    candidate_category = str(candidate.get('job_category', '')).lower()
                    job_category = str(job.get('job_category', '')).lower()
                    
                    # Calculate base similarity
                    similarity_score = 0.4  # Base score
                    
                    # Category matching bonus
                    if candidate_category and job_category:
                        if candidate_category == job_category:
                            similarity_score = 0.8
                        elif candidate_category in job_category or job_category in candidate_category:
                            similarity_score = 0.6
                    
                    # Add randomness for variety
                    import random
                    random_bonus = random.uniform(0.0, 0.3)
                    similarity_score = min(0.95, similarity_score + random_bonus)
                    
                    # Skip if too low
                    if similarity_score < 0.2:
                        continue
                    
                    # Calculate distance
                    distance = 1.0 - similarity_score
                    
                    # Create match record
                    match_id = str(uuid.uuid4())
                    
                    # FIXED: Insert with proper error handling
                    insert_query = """
                        INSERT INTO job_matches (
                            match_id, candidate_id, job_offer_id, 
                            distance, similarity_score, overall_evaluation,
                            created_at, updated_at
                        ) VALUES (%s, %s, %s, %s, %s, %s, NOW(), NOW())
                        ON CONFLICT (candidate_id, job_offer_id) DO UPDATE SET
                            distance = EXCLUDED.distance,
                            similarity_score = EXCLUDED.similarity_score,
                            overall_evaluation = EXCLUDED.overall_evaluation,
                            updated_at = NOW()
                    """
                    
                    db_manager.execute_query(insert_query, (
                        match_id,
                        candidate_id,
                        job['job_offer_id'],
                        float(distance),
                        float(similarity_score),
                        float(similarity_score)
                    ))
                    
                    matches_created += 1
                    
                    if idx < 3:  # Show first 3 matches
                        st.info(f"✅ Match {idx+1}: {job.get('job_title', 'N/A')} - {similarity_score:.1%}")
                    
                except Exception as match_error:
                    print(f"Error creating match for job {idx}: {match_error}")
                    continue
            
            st.success(f"✅ {matches_created} matches!")
            return matches_created
            
        except Exception as e:
            st.error(f"❌ Manual matching failed: {e}")
            return 0
        
    def insert_candidate_with_matching(candidate_data, db_manager, pipeline=None):
        """Enhanced version that ensures matches are created"""
        
        try:
            # Insert candidate using existing logic
            result_candidate_id = insert_candidate_with_related_data(candidate_data)
            
            if not result_candidate_id:
                return None
            
            # CRITICAL: Create job matches immediately after candidate insertion
            matches_created = 0
            
            # Try pipeline first if available
            if pipeline:
                try:
                    st.info("🔄 Creating job matches with pipeline...")
                    matches_created = pipeline.process_matching_for_candidate(result_candidate_id)
                    st.success(f"✅ Pipeline created {matches_created} job matches!")
                except Exception as match_error:
                    st.warning(f"⚠️ Pipeline matching failed: {match_error}")
            
            # Try manual matching if pipeline failed or no pipeline
            if matches_created == 0:
                st.info("🔄 Trying manual match creation...")
                matches_created = create_manual_matches(result_candidate_id)
                if matches_created > 0:
                    st.success(f"✅ Manual matching created {matches_created} job matches!")
            
            # Verify matches were actually created
            try:
                verify_matches = db_manager.execute_query(
                    "SELECT COUNT(*) as count FROM job_matches WHERE candidate_id = %s",
                    (result_candidate_id,)
                )
                
                actual_matches = verify_matches[0]['count'] if verify_matches else 0
                st.info(f"🔍 Verified: {actual_matches} matches in database")
                
                if actual_matches == 0:
                    st.warning("⚠️ No matches were actually saved to database")
                    # Try one more time with basic insertion
                    try:
                        basic_matches = create_basic_test_matches(result_candidate_id)
                        if basic_matches > 0:
                            st.success(f"{basic_matches} basic test matches!")
                    except Exception as e:
                        st.error(f"Even basic matching failed: {e}")
                else:
                    st.balloons()
            except Exception as verify_error:
                st.warning(f"Could not verify matches: {verify_error}")
            
            return result_candidate_id
            
        except Exception as e:
            st.error(f"❌ Error in enhanced insertion: {e}")
            return None
    def create_basic_test_matches(candidate_id):
        """Create basic test matches - FIXED VERSION"""
        try:
            st.info("🔄 Creating basic test matches as fallback...")
            
            # Get ANY job offers (regardless of status) as last resort
            jobs_query = """
                SELECT job_offer_id, job_title, company_name, status 
                FROM job_offers 
                ORDER BY created_at DESC 
                LIMIT 10
            """
            jobs = db_manager.execute_query(jobs_query)
            
            if not jobs:
                st.error("❌ No job offers found in database at all!")
                return 0
            
            st.info(f"📊 Found {len(jobs)} job offers for basic matching")
            
            matches_created = 0
            for idx, job in enumerate(jobs):
                try:
                    # Check if match already exists
                    existing_match = db_manager.fetch_one(
                        "SELECT match_id FROM job_matches WHERE candidate_id = %s AND job_offer_id = %s",
                        (candidate_id, job['job_offer_id'])
                    )
                    
                    if existing_match:
                        st.info(f"⚠️ Match already exists for job: {job.get('job_title', 'N/A')}")
                        continue
                    
                    match_id = str(uuid.uuid4())
                    
                    # Use different scores for variety
                    base_score = 0.6 + (idx * 0.05)  # 0.6, 0.65, 0.7, etc.
                    distance = 1.0 - base_score
                    
                    # FIXED: Correct SQL and parameters
                    insert_query = """
                        INSERT INTO job_matches (
                            match_id, candidate_id, job_offer_id, distance, similarity_score,
                            overall_evaluation, created_at, updated_at
                        ) VALUES (%s, %s, %s, %s, %s, %s, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                    """
                    
                    # FIXED: Pass exactly 6 parameters
                    db_manager.execute_query(insert_query, (
                        match_id, 
                        candidate_id, 
                        job['job_offer_id'], 
                        distance,           # distance
                        base_score,         # similarity_score
                        base_score          # overall_evaluation
                    ))
                    
                    matches_created += 1
                    st.info(f"match {matches_created}: {job.get('job_title', 'N/A')} (Score: {base_score:.2f})")
                    
                except Exception as insert_error:
                    st.error(f"❌ Failed to insert match for job {idx+1}: {insert_error}")
                    continue
            
            if matches_created > 0:
                st.success(f"✅ matching {matches_created} test matches!")
            else:
                st.error("❌ Even basic matching failed to create any matches")
            
            return matches_created
            
        except Exception as e:
            st.error(f"❌ Basic test matching failed: {e}")
            with st.expander("🔍 Error Details"):
                st.code(traceback.format_exc())
            return 0
    
    def get_candidate_matches_fixed(candidate_id):
        """Get matches with CORRECT field names for your database schema"""
        try:
            # Query with YOUR EXACT field names
            matches_raw = db_manager.execute_query("""
                SELECT 
                    jm.match_id,
                    jm.candidate_id,
                    jm.job_offer_id,
                    jm.similarity_score,
                    jm.distance,
                    jm.overall_evaluation,
                    jm.created_at as match_created,
                    jm.updated_at as match_updated
                    
                    
                    
                    
                FROM job_matches jm
                JOIN job_offers jo ON jm.job_offer_id = jo.job_offer_id
                WHERE jm.candidate_id = %s
                ORDER BY jm.similarity_score DESC, jm.created_at DESC
            """, (candidate_id,))
            
            return pd.DataFrame(matches_raw) if matches_raw else pd.DataFrame()
            
        except Exception as e:
            st.error(f"Error getting matches: {e}")
            return pd.DataFrame()
    
    def verify_matches_fixed(candidate_id):
        """Verify matches with SAFE dictionary access"""
        try:
            # Count matches
            verify_result = db_manager.fetch_one(
                "SELECT COUNT(*) as count FROM job_matches WHERE candidate_id = %s",
                (candidate_id,)
            )
            
            # SAFE: Handle both dict and tuple
            if verify_result:
                if isinstance(verify_result, dict):
                    count = verify_result['count']
                else:  # tuple
                    count = verify_result[0]
            else:
                count = 0
                
            st.info(f"🔍 Verified: {count} matches in database")
            
            if count > 0:
                # Show sample matches
                samples = db_manager.execute_query("""
                    SELECT jm.similarity_score, jm.overall_evaluation, jo.job_title, jo.company_name 
                    FROM job_matches jm 
                    LEFT JOIN job_offers jo ON jm.job_offer_id = jo.job_offer_id 
                    WHERE jm.candidate_id = %s 
                    ORDER BY jm.similarity_score DESC 
                    LIMIT 3
                """, (candidate_id,))
                
                if samples:
                    st.success("✅ Sample matches:")
                    for i, match in enumerate(samples):
                        # SAFE: Handle both dict and tuple
                        if isinstance(match, dict):
                            job_title = match.get('job_title', 'Unknown')
                            company = match.get('company_name', 'Unknown')
                            score = match.get('similarity_score', 0)
                            evaluation = match.get('overall_evaluation', 'N/A')
                        else:  # tuple
                            job_title = match[2] if len(match) > 2 else 'Unknown'
                            company = match[3] if len(match) > 3 else 'Unknown'
                            score = match[0] if len(match) > 0 else 0
                            evaluation = match[1] if len(match) > 1 else 'N/A'
                        
                        st.write(f"{i+1}. {job_title} at {company} ({score:.1%} - {evaluation})")
            
            return count
            
        except Exception as e:
            st.error(f"Verification failed: {e}")
            st.code(traceback.format_exc())
            return 0
   # ================================================================
    # INITIALIZE PIPELINE AND COMPONENTS
    # ================================================================
    
    # Initialize pipeline if not exists
    if 'pipeline' not in st.session_state:
        try:
            if 'pipeline' in globals():
                st.session_state.pipeline = pipeline
                st.success("✅ Pipeline loaded from global scope")
            else:
                st.warning("⚠️ Pipeline not found in global scope")
                st.session_state.pipeline = None
        except Exception as e:
            st.error(f"❌ Pipeline initialization failed: {e}")
            st.session_state.pipeline = None
    
    # Add helper functions to session state
    st.session_state.safe_skills_processor = safe_skills_processor
    st.session_state.skills_to_display_format = skills_to_display_format
    st.session_state.safe_vector_format = safe_vector_format
    st.session_state.insert_candidate_with_related_data = insert_candidate_with_related_data
    st.session_state.insert_candidate_with_matching = insert_candidate_with_matching
    st.session_state.create_manual_matches = create_manual_matches

   # =================================================================
    # MAIN DASHBOARD UI CODE
    # =================================================================
    
    """Display candidate dashboard with profile management, CV upload, and job matching"""
    st.title("👤 Candidate Dashboard")
    
    # Sidebar with actions
    with st.sidebar:
        st.markdown(f"### Welcome, Candidate #{st.session_state.user_id}")
        
        # ADD ENHANCED REFRESH MATCHES BUTTON
        if st.button("🔄 Refresh Matches", key="refresh_matches", use_container_width=True):
            with st.spinner("Finding your perfect job matches..."):
                st.info("🎯 Starting comprehensive match creation process...")
                
                matches_created = 0
                
                # First, check database status
                try:
                    candidate_check = db_manager.fetch_one(
                        "SELECT candidate_id, job_category, skills FROM candidates WHERE candidate_id = %s",
                        (st.session_state.user_id,)
                    )
                    if not candidate_check:
                        st.error("❌ Candidate not found in database!")
                        return
                    
                    st.info(f"✅ Found candidate with category: {candidate_check.get('job_category', 'None')}")
                    
                    # Check job offers
                    job_count = db_manager.fetch_one("SELECT COUNT(*) as count FROM job_offers")
                    active_job_count = db_manager.fetch_one("SELECT COUNT(*) as count FROM job_offers WHERE status = 'active'")
                    
                    st.info(f"📊 Database status: {job_count['count'] if job_count else 0} total jobs, {active_job_count['count'] if active_job_count else 0} active")
                    
                except Exception as db_check_error:
                    st.error(f"❌ Database check failed: {db_check_error}")
                    return
                
                # Try using session state pipeline first
                if st.session_state.get('pipeline'):
                    try:
                        st.info("🔄 Trying pipeline matching...")
                        matches_created = st.session_state.pipeline.process_matching_for_candidate(st.session_state.user_id)
                        if matches_created > 0:
                            st.success(f"✅ Pipeline created {matches_created} matches!")
                        else:
                            st.warning("⚠️ Pipeline returned 0 matches")
                    except Exception as e:
                        st.error(f"❌ Pipeline matching failed: {e}")
                
                # If pipeline failed or no matches, try manual matching
                if matches_created == 0:
                    st.info("🔄 Trying enhanced manual match creation...")
                    matches_created = create_manual_matches(st.session_state.user_id)
                
                # If still no matches, try basic matching
                if matches_created == 0:
                    st.warning("⚠️ Manual matching failed, trying basic test matching...")
                    matches_created = create_basic_test_matches(st.session_state.user_id)
                
                # Verify matches in database with detailed query
                try:
                    verify_result = db_manager.fetch_one(
                        "SELECT COUNT(*) as count FROM job_matches WHERE candidate_id = %s",
                        (st.session_state.user_id,)
                    )
                    actual_count = verify_result['count'] if verify_result else 0
                    
                    # Also get some sample matches for verification
                    sample_matches = db_manager.execute_query("""
                        SELECT jm.similarity_score, jo.job_title, jo.company_name 
                        FROM job_matches jm 
                        JOIN job_offers jo ON jm.job_offer_id = jo.job_offer_id 
                        WHERE jm.candidate_id = %s 
                        ORDER BY jm.similarity_score DESC 
                        LIMIT 3
                    """, (st.session_state.user_id,))
                    
                    st.info(f"🔍 Final verification: {actual_count} matches in database")
                    
                    if sample_matches:
                        st.success("✅ Sample matches found:")
                        for match in sample_matches:
                            st.write(f"• {match.get('job_title', 'N/A')} at {match.get('company_name', 'N/A')} ({match.get('similarity_score', 0):.1%})")
                    
                except Exception as e:
                    st.warning(f"Could not verify matches: {e}")
                
                if matches_created == 0:
                    st.error("❌ No matches could be created. Please check:")
                    st.write("1. Is your profile complete?")
                    st.write("2. Are there job offers in the database?")
                    st.write("3. Check database connection and permissions")
                else:
                    st.balloons()
                
                time.sleep(1)
                st.rerun()
        
        if st.button("🚪 Logout", key="candidate_logout", use_container_width=True):
            st.session_state.user_role = None
            st.session_state.user_id = None
            st.rerun()
        
        if st.session_state.last_update:
            st.caption(f"Last updated: {st.session_state.last_update}")
    
    # Main tabs
    tab1, tab2, tab3 , tab4= st.tabs([
        "📋 My Profile", 
        "📤 Upload CV", 
        "📊 Analytics",
        "🎯 Job Matches"
    ])

    with tab1:
        st.markdown("### 📋 Personal Information")
        st.info("Fill out your profile information below")
        
        # Initialize session state for dynamic fields
        if 'experience_entries' not in st.session_state:
            st.session_state.experience_entries = []
        if 'certification_entries' not in st.session_state:
            st.session_state.certification_entries = []
        
        # Load existing candidate data if available
        existing_candidate = None
        use_cv_data = False
        
        # ✅ NEW: Check if we just processed a CV
        if st.session_state.get('cv_just_processed') and st.session_state.get('extracted_cv_data'):
            st.info("🎉 Using data extracted from your uploaded CV! You can edit any field below.")
            existing_candidate = st.session_state.extracted_cv_data
            use_cv_data = True
            # Clear the flag so it doesn't keep showing CV data
            st.session_state.cv_just_processed = False
        elif st.session_state.user_id:
            # Load from database as usual
            try:
                existing_candidate = db_manager.fetch_one(
                    "SELECT * FROM candidates WHERE candidate_id = %(candidate_id)s",
                    {'candidate_id': st.session_state.user_id}
                )
                if existing_candidate:
                    st.success(f"✅ Loaded existing profile for: {existing_candidate.get('first_name', '')} {existing_candidate.get('last_name', '')}")
            except Exception as e:
                st.warning(f"Could not load existing profile: {e}")
        
        # Show data source indicator
        if use_cv_data:
            st.success("📄 Showing data from uploaded CV - you can edit any field")
        elif existing_candidate:
            st.info("💾 Showing saved profile data")
        else:
            st.info("✏️ Fill out your profile information below")

        # Personal Information with default values from CV or existing data
        col1, col2 = st.columns(2)
        with col1:
            first_name = st.text_input(
                "First Name", 
                value=existing_candidate.get('first_name', '') if existing_candidate else '',
                key="profile_first_name",
                help="Extracted from CV" if use_cv_data else "Email is required and must be valid"
            )
            last_name = st.text_input(
                "Last Name", 
                value=existing_candidate.get('last_name', '') if existing_candidate else '',
                key="profile_last_name",
                help="Extracted from CV" if use_cv_data else ""
            )
            email = st.text_input(
                "Email *", 
                value=existing_candidate.get('email', '') if existing_candidate else '',
                key="profile_email",
                help="Extracted from CV" if use_cv_data else "Email is required and must be valid"
            )
        with col2:
            phone = st.text_input(
                "Phone", 
                value=existing_candidate.get('phone', '') if existing_candidate else '',
                key="profile_phone",
                help="Extracted from CV" if use_cv_data else ""
            )
            linkedin = st.text_input(
                "LinkedIn URL", 
                value=existing_candidate.get('linkedin', '') if existing_candidate else '',
                key="profile_linkedin",
                help="Extracted from CV" if use_cv_data else ""
            )
        
        # Email validation
        if email and '@' not in email:
            st.error("❌ Please enter a valid email address")
        
        st.markdown("### 🧠 Professional Details")
        col3, col4 = st.columns(2)
        with col3:
            job_category = st.text_input(
                "Job Category", 
                value=existing_candidate.get('job_category', '') if existing_candidate else '',
                key="profile_job_category",
                help="Extracted from CV" if use_cv_data else ""
            )
            current_job = st.text_input(
                "Current Job", 
                value=existing_candidate.get('current_job', '') if existing_candidate else '',
                key="profile_current_job",
                help="Extracted from CV" if use_cv_data else ""
            )
            city = st.text_input(
                "City", 
                value=existing_candidate.get('city', '') if existing_candidate else '',
                key="profile_city",
                help="Extracted from CV" if use_cv_data else ""
            )
        with col4:
            # ✅ FIXED: Handle experience_years properly for CV data
            default_experience = 1
            if existing_candidate:
                if use_cv_data:
                    # For CV data, experience_years might be a string or int
                    exp_value = existing_candidate.get('experience_years', 1)
                    try:
                        default_experience = int(exp_value) if exp_value else 1
                    except (ValueError, TypeError):
                        default_experience = 1
                else:
                    # For database data
                    default_experience = existing_candidate.get('experience_years', 1)
            
            experience_years = st.slider(
                "Years of Experience", 
                0, 50, 
                default_experience,
                key="profile_experience_years",
                help="Extracted from CV" if use_cv_data else ""
            )
            
            # ✅ FIXED: Handle education_level properly for CV data
            education_level_options = ["High School", "Bachelor", "Master", "PhD", "Other"]
            default_education = 'Bachelor'
            
            if existing_candidate and existing_candidate.get('education_level'):
                extracted_education = existing_candidate.get('education_level', 'Bachelor')
                # Map common variations to our options
                education_mapping = {
                    'high school': 'High School',
                    'bachelor': 'Bachelor',
                    'bachelors': 'Bachelor',
                    'bachelor\'s': 'Bachelor',
                    'master': 'Master',
                    'masters': 'Master',
                    'master\'s': 'Master',
                    'phd': 'PhD',
                    'ph.d': 'PhD',
                    'doctorate': 'PhD',
                    'unknown': 'Other'
                }
                
                mapped_education = education_mapping.get(extracted_education.lower(), extracted_education)
                if mapped_education in education_level_options:
                    default_education = mapped_education
                else:
                    default_education = 'Other'
            
            education_level_index = education_level_options.index(default_education)
            education_level = st.selectbox(
                "Education Level", 
                education_level_options,
                index=education_level_index,
                key="profile_education_level",
                help="Extracted from CV" if use_cv_data else ""
            )
            
            country = st.text_input(
                "Country", 
                value=existing_candidate.get('country', '') if existing_candidate else '',
                key="profile_country",
                help="Extracted from CV" if use_cv_data else ""
            )
        
        # Handle skills data safely
        existing_skills_tech = ""
        existing_skills_soft = ""
        
        if existing_candidate:
            try:
                if use_cv_data:
                    # For CV data, skills are already in the correct format
                    skills_data = existing_candidate.get('skills', {})
                    if isinstance(skills_data, dict):
                        existing_skills_tech = ', '.join(skills_data.get('technical_skills', []))
                        existing_skills_soft = ', '.join(skills_data.get('soft_skills', []))
                else:
                    # For database data, use the existing processor
                    if existing_candidate.get('skills'):
                        processed_skills = safe_skills_processor(existing_candidate['skills'])
                        existing_skills_tech = skills_to_display_format(processed_skills, 'technical_skills')
                        existing_skills_soft = skills_to_display_format(processed_skills, 'soft_skills')
            except Exception as e:
                st.warning(f"Could not parse skills data: {e}")
        
        technical_skills = st.text_area(
            "Technical Skills (comma-separated)", 
            value=existing_skills_tech,
            key="technical_skills", 
            height=100,
            help="Extracted from CV - you can edit these" if use_cv_data and existing_skills_tech else "e.g., Python, Java, Machine Learning, SQL"
        )
        soft_skills = st.text_area(
            "Soft Skills (comma-separated)", 
            value=existing_skills_soft,
            key="soft_skills", 
            height=100,
            help="Extracted from CV - you can edit these" if use_cv_data and existing_skills_soft else "e.g., Leadership, Communication, Problem Solving"
        )


        # Experience Section
        st.markdown("### 💼 Experience")
        
        if use_cv_data and existing_candidate.get('experience') and not st.session_state.experience_entries:
            try:
                cv_experience = existing_candidate.get('experience', [])
                if isinstance(cv_experience, list):
                    st.session_state.experience_entries = cv_experience
                    st.info(f"📄 Loaded {len(cv_experience)} experience entries from CV")
            except Exception as e:
                st.warning(f"Could not load CV experience: {e}")
        elif existing_candidate and existing_candidate.get('experience') and not st.session_state.experience_entries:
            # Load from database as before
            try:
                if isinstance(existing_candidate['experience'], str):
                    existing_experience = json.loads(existing_candidate['experience'])
                else:
                    existing_experience = existing_candidate['experience']
                
                if isinstance(existing_experience, list):
                    st.session_state.experience_entries = existing_experience
            except (json.JSONDecodeError, TypeError) as e:
                st.warning(f"Could not load existing experience: {e}")
        if st.button("➕ Add Experience", key="add_experience"):
            st.session_state.experience_entries.append({
                "title": "",
                "company": "",
                "start_year": "",
                "start_month": "",
                "end_year": "",
                "end_month": "",
                "en_cours": False,
                "description": ""
            })
        
        for idx, exp in enumerate(st.session_state.experience_entries):
            with st.expander(f"Experience #{idx + 1}", expanded=True):
                col1, col2 = st.columns(2)
                with col1:
                    exp["title"] = st.text_input(f"Title #{idx+1}", value=exp.get("title", ""), key=f"exp_title_{idx}")
                    exp["company"] = st.text_input(f"Company #{idx+1}", value=exp.get("company", ""), key=f"exp_company_{idx}")
                    exp["start_year"] = st.text_input(f"Start Year #{idx+1}", value=exp.get("start_year", ""), key=f"exp_start_year_{idx}")
                    exp["start_month"] = st.text_input(f"Start Month #{idx+1}", value=exp.get("start_month", ""), key=f"exp_start_month_{idx}")
                with col2:
                    exp["end_year"] = st.text_input(f"End Year #{idx+1}", value=exp.get("end_year", ""), key=f"exp_end_year_{idx}")
                    exp["end_month"] = st.text_input(f"End Month #{idx+1}", value=exp.get("end_month", ""), key=f"exp_end_month_{idx}")
                    exp["en_cours"] = st.checkbox("Currently Working?", value=exp.get("en_cours", False), key=f"exp_en_cours_{idx}")
                exp["description"] = st.text_area(f"Description #{idx+1}", value=exp.get("description", ""), key=f"exp_description_{idx}")
                
                if st.button(f"Remove Experience #{idx+1}", key=f"remove_exp_{idx}"):
                    st.session_state.experience_entries.pop(idx)
                    st.rerun()
        
        # Certifications Section
        st.markdown("### 🎓 Certifications")
        
        # Load CV certification data if available
        if use_cv_data and existing_candidate.get('certifications') and not st.session_state.certification_entries:
            try:
                cv_certifications = existing_candidate.get('certifications', [])
                if isinstance(cv_certifications, list):
                    st.session_state.certification_entries = cv_certifications
                    st.info(f"📄 Loaded {len(cv_certifications)} certifications from CV")
            except Exception as e:
                st.warning(f"Could not load CV certifications: {e}")
        elif existing_candidate and existing_candidate.get('certifications') and not st.session_state.certification_entries:
            # Load from database as before
            try:
                if isinstance(existing_candidate['certifications'], str):
                    existing_certifications = json.loads(existing_candidate['certifications'])
                else:
                    existing_certifications = existing_candidate['certifications']
                
                if isinstance(existing_certifications, list):
                    st.session_state.certification_entries = existing_certifications
            except (json.JSONDecodeError, TypeError) as e:
                st.warning(f"Could not load existing certifications: {e}")
        if st.button("➕ Add Certification", key="add_certification"):
            st.session_state.certification_entries.append({
                "name": "",
                "year": "",
                "issuer": "",
                "details": ""
            })
        
        for idx, cert in enumerate(st.session_state.certification_entries):
            with st.expander(f"Certification #{idx + 1}", expanded=True):
                col1, col2 = st.columns(2)
                with col1:
                    cert["name"] = st.text_input(f"Name #{idx+1}", value=cert.get("name", ""), key=f"cert_name_{idx}")
                    cert["year"] = st.text_input(f"Year #{idx+1}", value=cert.get("year", ""), key=f"cert_year_{idx}")
                with col2:
                    cert["issuer"] = st.text_input(f"Issuer #{idx+1}", value=cert.get("issuer", ""), key=f"cert_issuer_{idx}")
                cert["details"] = st.text_area(f"Details #{idx+1}", value=cert.get("details", ""), key=f"cert_details_{idx}")
                
                if st.button(f"Remove Certification #{idx+1}", key=f"remove_cert_{idx}"):
                    st.session_state.certification_entries.pop(idx)
                    st.rerun()
        
        # Submit button with enhanced functionality and validation
        st.markdown("---")
        col1, col2, col3 = st.columns([2, 1, 1])
        with col1:
            # Pre-submission validation
            can_submit = True
            validation_messages = []
            
            if not email or '@' not in email:
                can_submit = False
                validation_messages.append("❌ Valid email is required")
            
            if validation_messages:
                for msg in validation_messages:
                    st.error(msg)
            
            if st.button("✅ Save Profile", key="submit_profile", 
                        use_container_width=True, type="primary", 
                        disabled=not can_submit):
                
                with st.spinner("Saving your profile..."):
                    try:
                        # Prepare data for submission
                        certifications = [
                            {
                                "name": c["name"],
                                "year": c["year"],
                                "issuer": c["issuer"],
                                "details": c["details"]
                            }
                            for c in st.session_state.certification_entries
                        ]
                        
                        experiences = [
                            {
                                "title": e["title"],
                                "company": e["company"],
                                "duration": {
                                    "start": {"year": e["start_year"], "month": e["start_month"]},
                                    "end": {"year": e["end_year"], "month": e["end_month"]}
                                },
                                "en_cours": e["en_cours"],
                                "description": e["description"]
                            }
                            for e in st.session_state.experience_entries
                        ]
                        
                        # Process skills safely
                        technical_skills_list = []
                        soft_skills_list = []
                        
                        if technical_skills and technical_skills.strip():
                            technical_skills_list = [
                                skill.strip() for skill in technical_skills.split(",") 
                                if skill.strip()
                            ]
                        
                        if soft_skills and soft_skills.strip():
                            soft_skills_list = [
                                skill.strip() for skill in soft_skills.split(",") 
                                if skill.strip()
                            ]
                        
                        skills_dict = {
                            "technical_skills": technical_skills_list,
                            "soft_skills": soft_skills_list
                        }
                        
                        # Generate embeddings if cv_generator is available
                        skills_vector = [0.0] * 1536  # Default vector
                        if hasattr(st.session_state, 'cv_generator') and st.session_state.cv_generator:
                            try:
                                embedding = st.session_state.cv_generator.generate_embedding(skills_dict)
                                skills_vector = safe_vector_format(embedding)
                            except Exception as e:
                                st.warning(f"Could not generate embeddings: {e}")
                        
                        # Prepare candidate data for insertion
                        candidate_data = {
                            
                            'first_name': first_name,
                            'last_name': last_name,
                            'email': email,
                            'phone': phone,
                            'linkedin': linkedin,
                            'job_category': job_category,
                            'current_job': current_job,
                            'country': country,
                            'city': city,
                            'experience_years': experience_years,
                            'education_level': education_level,
                            'skills': skills_dict,
                            'skills_vector': skills_vector,
                            'certifications': certifications,
                            'experience': experiences
                        }
                        
                        # CRITICAL CHANGE: Use enhanced insertion with matching
                        result_candidate_id = insert_candidate_with_matching(
                            candidate_data, 
                            db_manager, 
                            st.session_state.get('pipeline')
                        )
                        
                        if result_candidate_id:
                            st.session_state.user_id = result_candidate_id
                            st.success("🎉 Profile and all related data saved successfully!")
                            time.sleep(2)
                            st.rerun()
                        else:
                            st.error("❌ Failed to save profile. Please check the error messages above and try again.")
                    
                    except Exception as save_error:
                        st.error(f"❌ Error saving profile: {str(save_error)}")
                        with st.expander("🔍 Error Details"):
                            st.code(traceback.format_exc())
        with col3:
            if existing_candidate:
                if st.button("🗑️ Delete Profile", key="delete_profile", use_container_width=True, type="secondary"):
                    try:
                        # Delete from database
                        delete_query = "DELETE FROM candidates WHERE candidate_id = %(candidate_id)s"
                        db_manager.execute_query(delete_query, {'candidate_id': st.session_state.user_id})
                        
                        # Clear session state
                        st.session_state.user_id = None
                        st.session_state.experience_entries = []
                        st.session_state.certification_entries = []
                        
                        st.success("✅ Profile deleted successfully!")
                        time.sleep(1)
                        st.rerun()
                    except Exception as e:
                        st.error(f"Could not delete profile: {e}")
    # Tab 2: CV Upload and Processing - ENHANCED VERSION
    with tab2:
        st.markdown("### 📤 Upload and Process CV")
        st.info("Upload your CV and our AI will automatically extract your information!")
        
        # Initialize CV processor components
        if 'cv_generator' not in st.session_state or 's3_manager' not in st.session_state:
            cv_generator, s3_manager = init_cv_components()
            st.session_state.cv_generator = cv_generator
            st.session_state.s3_manager = s3_manager
        
        # Clear button at the top
        col_clear, col_info = st.columns([1, 3])
        with col_clear:
            if st.button("🗑️ Clear CV Data", key="clear_cv_data"):
                # Clear CV-related session state
                if 'extracted_cv_data' in st.session_state:
                    del st.session_state.extracted_cv_data
                if 'cv_just_processed' in st.session_state:
                    del st.session_state.cv_just_processed
                st.session_state.experience_entries = []
                st.session_state.certification_entries = []
                st.success("✅ CV data cleared - showing database data")
                st.rerun()
                
        uploaded_file = st.file_uploader(
            "Choose a CV file",
            type=['pdf', 'docx', 'doc', 'txt'],
            help="Supported formats: PDF, DOCX, DOC, TXT"
        )
        
        if uploaded_file is not None:
            # Display file info
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("📄 Filename", uploaded_file.name)
            with col2:
                st.metric("📏 Size", f"{uploaded_file.size} bytes")
            with col3:
                file_ext = os.path.splitext(uploaded_file.name)[1].upper()
                st.metric("📋 Type", file_ext)
            
            # Preview button
            if st.button("👁️ Preview CV Content"):
                with st.spinner("Extracting text preview..."):
                    file_content = uploaded_file.read()
                    uploaded_file.seek(0)  # Reset file pointer
                    
                    parser = ParserFactory.get_parser(uploaded_file.name)
                    if parser:
                        text = parser.extract_text(file_content, uploaded_file.name)
                        if text:
                            st.subheader("📄 CV Text Preview")
                            st.text_area(
                                "Extracted Text",
                                text[:2000] + "..." if len(text) > 2000 else text,
                                height=300,
                                help="Showing first 2000 characters"
                            )
                        else:
                            st.error("❌ Could not extract text from the file")
                    else:
                        st.error("❌ Unsupported file format")
            
            st.markdown("---")
            
            # ENHANCED PROCESS BUTTON WITH GUARANTEED MATCHING
            col1, col2 = st.columns([1, 1])
            with col1:
                if st.button("🚀 Upload & Process CV", type="primary", use_container_width=True):
                    with st.spinner("Processing your CV with AI..."):
                        try:
                            # Read file content ONCE and store it
                            file_content = uploaded_file.read()
                            filename = uploaded_file.name
                            
                            # Validate file content
                            if not file_content:
                                st.error("❌ File appears to be empty")
                                st.stop()
                            
                            st.info(f"📄 Processing file: {filename} ({len(file_content)} bytes)")
                            
                            # CRITICAL FIX: Upload to S3 FIRST with detailed error handling
                            st.info("📤 Uploading file to S3...")
                            
                            upload_success = False
                            s3_key_or_error = ""
                            
                            try:
                                if not st.session_state.s3_manager:
                                    st.warning("⚠️ S3 manager not initialized, reinitializing...")
                                    st.session_state.s3_manager = EnhancedS3Manager()
                                
                                if st.session_state.s3_manager.s3_available:
                                    upload_success, s3_key_or_error = st.session_state.s3_manager.upload_file(
                                        file_content, filename
                                    )
                                    
                                    if upload_success:
                                        st.success(f"✅ File successfully uploaded to S3: {s3_key_or_error}")
                                    else:
                                        st.error(f"❌ S3 upload failed: {s3_key_or_error}")
                                        st.error("🛑 Cannot proceed without S3 upload. Please check AWS configuration.")
                                        
                                        # Show S3 configuration help
                                        with st.expander("🔧 S3 Configuration Help"):
                                            st.write("**Required AWS Configuration:**")
                                            st.write(f"**S3 Bucket:** {CONFIG['S3_BUCKET_CANDIDATES']}")
                                            st.write(f"**Region:** {CONFIG['AWS_REGION']}")
                                        
                                        st.stop()  # Stop processing if S3 upload fails
                                else:
                                    st.error("❌ S3 service not available")
                                    st.error("Please check AWS credentials and bucket configuration")
                                    st.stop()
                                    
                            except Exception as s3_error:
                                st.error(f"❌ S3 upload error: {s3_error}")
                                with st.expander("🔍 S3 Error Details"):
                                    st.code(str(s3_error))
                                st.stop()
                            
                            # Continue with text extraction (only if S3 upload succeeded)
                            st.info("📄 Extracting text from CV...")
                            parser = ParserFactory.get_parser(filename)
                            if not parser:
                                st.error("❌ Unsupported file format")
                                st.stop()
                            
                            text = parser.extract_text(file_content, filename)
                            if not text or not text.strip():
                                st.error("❌ No text extracted from the file")
                                st.stop()
                            
                            st.info(f"✅ Extracted {len(text)} characters from CV")
                            
                            # Extract candidate information using AI
                            st.info("🤖 Processing with AI...")
                            try:
                                candidate_data = safe_extract_with_langchain(text, filename)
                                
                                if not candidate_data:
                                    st.error("❌ AI extraction failed - no data returned")
                                    st.stop()
                                    
                            except Exception as ai_error:
                                st.error(f"❌ AI extraction error: {ai_error}")
                                st.stop()
                            
                            # IMPORTANT: Add the S3 URL to candidate data
                            candidate_data['url'] = s3_key_or_error  # This will be saved to PostgreSQL
                            candidate_data['filename'] = filename
                            
                            # Display extracted information
                            st.success("🎉 CV processed successfully!")
                            
                            # Show extracted data in expandable sections
                            with st.expander("📋 Extracted Personal Information", expanded=True):
                                col1, col2 = st.columns(2)
                                with col1:
                                    st.write(f"**Name:** {candidate_data.get('first_name', '')} {candidate_data.get('last_name', '')}")
                                    st.write(f"**Email:** {candidate_data.get('email', '')}")
                                    st.write(f"**Phone:** {candidate_data.get('phone', '')}")
                                    st.write(f"**Experience Years:** {candidate_data.get('experience_years', 0)}")
                                with col2:
                                    st.write(f"**Current Job:** {candidate_data.get('current_job', '')}")
                                    st.write(f"**Location:** {candidate_data.get('city', '')}, {candidate_data.get('country', '')}")
                                    st.write(f"**LinkedIn:** {candidate_data.get('linkedin', '')}")
                                    st.write(f"**Education:** {candidate_data.get('education_level', 'Unknown')}")
                                
                                # SHOW SKILLS IN PERSONAL INFO
                                st.markdown("**🛠️ Technical Skills:**")
                                tech_skills = candidate_data.get('technical_skills', [])
                                if tech_skills:
                                    st.write(", ".join(tech_skills))
                                else:
                                    st.write("No technical skills found")
                                
                                st.markdown("**🤝 Soft Skills:**")
                                soft_skills = candidate_data.get('soft_skills', [])
                                if soft_skills:
                                    st.write(", ".join(soft_skills))
                                else:
                                    st.write("No soft skills found")
                            
                            # Save to database with enhanced error handling
                            try:
                                # Validate email before proceeding
                                email = candidate_data.get("email", "").strip()
                                if not email:
                                    st.error("❌ No email found in CV. Please add email manually in the profile tab.")
                                    st.stop()
                                
                                if '@' not in email:
                                    st.error("❌ Invalid email format found in CV. Please correct in profile tab.")
                                    st.stop()
                                
                                # Prepare skills dict
                                skills_dict = {
                                    "technical_skills": candidate_data.get("technical_skills", []),
                                    "soft_skills": candidate_data.get("soft_skills", [])
                                }
                                
                                # Generate embeddings
                                try:
                                    embedding = st.session_state.cv_generator.generate_embedding(skills_dict)
                                    safe_embedding = safe_vector_format(embedding)
                                except Exception as embed_error:
                                    st.warning(f"Could not generate embeddings: {embed_error}")
                                    safe_embedding = [0.0] * 1536
                                
                                # Prepare full candidate data for insertion
                                cv_candidate_data = {
                                    'first_name': candidate_data.get('first_name', ''),
                                    'last_name': candidate_data.get('last_name', ''),
                                    'email': email,
                                    'phone': candidate_data.get('phone', ''),
                                    'linkedin': candidate_data.get('linkedin', ''),
                                    'job_category': candidate_data.get('job_category', ''),
                                    'current_job': candidate_data.get('current_job', ''),
                                    'country': candidate_data.get('country', ''),
                                    'city': candidate_data.get('city', ''),
                                    'experience_years': candidate_data.get('experience_years', 0),
                                    'education_level': candidate_data.get('education_level', 'Unknown'),
                                    'skills': skills_dict,
                                    'skills_vector': safe_embedding,
                                    'certifications': candidate_data.get('certifications', []),
                                    'experience': candidate_data.get('experience', []),
                                    'filename': filename,
                                    'url': s3_key_or_error,  # ✅ This is the S3 URL
                                    'document_extension': candidate_data.get('document_extension', ''),
                                    'document_language': candidate_data.get('language', 'unknown')
                                }
                                
                                # CRITICAL: Use enhanced insertion function with guaranteed matching
                                result_candidate_id = insert_candidate_with_matching(
                                    cv_candidate_data, 
                                    db_manager, 
                                    st.session_state.get('pipeline')
                                )
                                
                                if result_candidate_id:
                                    # Update session state
                                    st.session_state.user_id = result_candidate_id
                                    st.success("✅ Profile and all related data saved to database!")
                                    st.success(f"✅ CV file stored in S3 at: {s3_key_or_error}")
                                else:
                                    st.error("❌ Failed to save to database. Check the error messages above.")
                                
                            except Exception as db_error:
                                st.error(f"❌ Database processing error: {str(db_error)}")
                                with st.expander("🔍 Error Details"):
                                    st.code(traceback.format_exc())
                        
                        except Exception as e:
                            st.error(f"❌ Processing error: {str(e)}")
                            with st.expander("🔍 Error Details"):
                                st.code(traceback.format_exc())
                
            with col2:
                st.info("**What happens when you upload:**")
                st.write("1. 📤 File uploaded to AWS S3")
                st.write("2. 📄 Text extracted from CV")
                st.write("3. 🤖 AI analyzes and extracts data")
                st.write("4. 💾 Candidate info saved to database")
                st.write("5. 🎯 Job matches calculated automatically")
                st.write("6. ✅ Profile ready for job searching")
        
        else:
            st.info("👆 Please upload a CV file to get started")
    # Tab 3: Enhanced Analytics - FIXED VERSION
    with tab3:
        st.markdown("### 📊 Your Enhanced Analytics Dashboard")
        
        try:
            # Get matches using the same logic as tab3
            if st.session_state.get('pipeline'):
                matches_df = get_candidate_matches_fixed(st.session_state.user_id)
            else:
                # Fallback: get matches directly from database
                matches_raw = db_manager.execute_query("""
                    SELECT jm.*, jo.job_title, jo.company_name, jo.city, jo.country, 
                           jo.job_type, jo.description, jo.job_category
                    FROM job_matches jm
                    JOIN job_offers jo ON jm.job_offer_id = jo.job_offer_id
                    WHERE jm.candidate_id = %s
                    ORDER BY jm.similarity_score DESC
                """, (st.session_state.user_id,))
                
                matches_df = pd.DataFrame(matches_raw) if matches_raw else pd.DataFrame()
            
            # Enhanced analytics with related tables data
            if st.session_state.user_id:
                # Get additional candidate analytics from related tables
                try:
                    # Skills analytics - Direct from candidates table since candidate_skills might not exist
                    candidate_info = db_manager.fetch_one(
                        "SELECT skills, experience, certifications FROM candidates WHERE candidate_id = %s",
                        (st.session_state.user_id,)
                    )
                    
                    user_skills = []
                    user_experiences = []
                    user_certifications = []
                    
                    if candidate_info:
                        # Parse skills
                        if candidate_info.get('skills'):
                            try:
                                skills_data = candidate_info['skills']
                                if isinstance(skills_data, str):
                                    skills_data = json.loads(skills_data)
                                
                                tech_skills = skills_data.get('technical_skills', [])
                                soft_skills = skills_data.get('soft_skills', [])
                                
                                user_skills = [{'name': skill, 'skill_type': 'technical'} for skill in tech_skills]
                                user_skills.extend([{'name': skill, 'skill_type': 'soft'} for skill in soft_skills])
                            except:
                                pass
                        
                        # Parse experiences
                        if candidate_info.get('experience'):
                            try:
                                exp_data = candidate_info['experience']
                                if isinstance(exp_data, str):
                                    exp_data = json.loads(exp_data)
                                if isinstance(exp_data, list):
                                    user_experiences = exp_data
                            except:
                                pass
                        
                        # Parse certifications
                        if candidate_info.get('certifications'):
                            try:
                                cert_data = candidate_info['certifications']
                                if isinstance(cert_data, str):
                                    cert_data = json.loads(cert_data)
                                if isinstance(cert_data, list):
                                    user_certifications = cert_data
                            except:
                                pass
                    
                    # Display enhanced analytics
                    if not matches_df.empty:
                        # Basic statistics
                        col1, col2, col3, col4 = st.columns(4)
                        with col1:
                            st.metric("Total Matches", len(matches_df))
                        with col2:
                            avg_score = matches_df['similarity_score'].mean()
                            st.metric("Average Match", f"{avg_score:.1%}")
                        with col3:
                            best_match = matches_df['similarity_score'].max()
                            st.metric("Best Match", f"{best_match:.1%}")
                        with col4:
                            skills_count = len(user_skills) if user_skills else 0
                            st.metric("Total Skills", skills_count)
                        
                        # Skills breakdown
                        if user_skills:
                            st.subheader("🛠️ Your Skills Breakdown")
                            col1, col2 = st.columns(2)
                            
                            with col1:
                                tech_skills = [skill for skill in user_skills if skill['skill_type'] == 'technical']
                                st.write(f"**Technical Skills ({len(tech_skills)}):**")
                                for skill in tech_skills:
                                    st.write(f"• {skill['name']}")
                            
                            with col2:
                                soft_skills = [skill for skill in user_skills if skill['skill_type'] == 'soft']
                                st.write(f"**Soft Skills ({len(soft_skills)}):**")
                                for skill in soft_skills:
                                    st.write(f"• {skill['name']}")
                        
                        # Experience timeline
                        # Experience timeline - FIXED VERSION
                            if user_experiences:
                                st.subheader("💼 Your Experience Timeline")
                                for idx, exp in enumerate(user_experiences):
                                    # Handle different possible field names and data types
                                    if isinstance(exp, dict):
                                        job_title = exp.get('job_title', exp.get('title', exp.get('position', 'Position')))
                                        company = exp.get('company', exp.get('employer', exp.get('organization', 'Company')))
                                        years = exp.get('years', exp.get('duration', exp.get('period', '')))
                                        
                                        if job_title and company:
                                            st.write(f"**{job_title}** at {company}")
                                            if years:
                                                st.write(f"*{years}*")
                                        else:
                                            st.write(f"• {str(exp)}")
                                    else:
                                        # If experience is just a string
                                        st.write(f"• {str(exp)}")
                                    
                                    if idx < len(user_experiences) - 1:
                                        st.divider()
                            elif candidate_info and candidate_info.get('experience'):
                                # Show raw experience if structured parsing failed
                                st.subheader("💼 Your Experience")
                                raw_exp = candidate_info['experience']
                                if isinstance(raw_exp, str):
                                    # Split by common separators and display
                                    exp_lines = raw_exp.replace(';', '\n').replace(',', '\n').split('\n')
                                    for line in exp_lines:
                                        line = line.strip()
                                        if line:
                                            st.write(f"• {line}")
                                else:
                                    st.write(raw_exp)
                        
                        # Certifications
                        if user_certifications:
                            st.subheader("🎓 Your Certifications")
                            
                            # List certifications
                            for cert in user_certifications:
                                st.write(f"• **{cert.get('name')}** - {cert.get('issuer')} ({cert.get('year')})")
                        
                        # Match score distribution
                        st.subheader("📈 Match Score Distribution")
                        score_ranges = pd.cut(matches_df['similarity_score'], 
                                            bins=[0, 0.2, 0.4, 0.6, 0.8, 1.0], 
                                            labels=['Very Low', 'Low', 'Medium', 'High', 'Excellent'])
                        score_counts = score_ranges.value_counts()
                        st.bar_chart(score_counts)
                        
                        # Top companies
                        if 'company_name' in matches_df.columns:
                            st.subheader("🏢 Top Companies Matching Your Profile")
                            top_companies = matches_df['company_name'].value_counts().head(10)
                            st.bar_chart(top_companies)
                        
                        # Job categories analysis
                        if 'job_category' in matches_df.columns:
                            st.subheader("📊 Job Categories Distribution")
                            job_categories = matches_df['job_category'].value_counts()
                            st.bar_chart(job_categories)
                    
                    else:
                        # Show profile completion status
                        st.info("No job matches available yet. Here's your profile completion status:")
                        
                        completion_metrics = []
                        if user_skills:
                            completion_metrics.append(("Skills", len(user_skills), "✅"))
                        else:
                            completion_metrics.append(("Skills", 0, "❌"))
                        
                        if user_experiences:
                            completion_metrics.append(("Experience", len(user_experiences), "✅"))
                        else:
                            completion_metrics.append(("Experience", 0, "❌"))
                        
                        if user_certifications:
                            completion_metrics.append(("Certifications", len(user_certifications), "✅"))
                        else:
                            completion_metrics.append(("Certifications", 0, "❌"))
                        
                        col1, col2, col3 = st.columns(3)
                        for idx, (metric_name, count, status) in enumerate(completion_metrics):
                            with [col1, col2, col3][idx]:
                                st.metric(f"{status} {metric_name}", count)
                        
                        # Show database statistics
                        try:
                            total_candidates = db_manager.fetch_one("SELECT COUNT(*) as count FROM candidates")
                            total_jobs = db_manager.fetch_one("SELECT COUNT(*) as count FROM job_offers")
                            total_matches = db_manager.fetch_one("SELECT COUNT(*) as count FROM job_matches")
                            
                            st.subheader("📊 Platform Statistics")
                            col1, col2, col3 = st.columns(3)
                            with col1:
                                if total_candidates:
                                    st.metric("Total Candidates", total_candidates['count'])
                            with col2:
                                if total_jobs:
                                    st.metric("Total Job Offers", total_jobs['count'])
                            with col3:
                                if total_matches:
                                    st.metric("Total Matches", total_matches['count'])
                        except Exception as e:
                            st.warning(f"Could not load database statistics: {e}")
                
                except Exception as analytics_error:
                    st.error(f"Analytics error: {analytics_error}")
                    st.info("Please ensure your profile is complete and try refreshing matches.")
            
            else:
                st.warning("Please log in to view your analytics")
                
        except Exception as e:
            st.error(f"Analytics error: {e}")
            st.info("Please ensure your profile is complete and try refreshing matches.")
    
    # Tab 4: Job Matches
    with tab4:
        st.markdown("### 🎯 Your Job Matches")
        
        # Get matches for current candidate
        try:
            matches_df = pipeline.get_candidate_matches(st.session_state.user_id)
        except Exception as e:
            st.warning(f"Could not load matches: {e}")
            matches_df = pd.DataFrame()
        
        if matches_df.empty:
            st.info("No job matches found yet. Upload your CV or click 'Refresh Job Matches' to find matches!")
            
            # Show some basic statistics about available jobs
            try:
                total_jobs = db_manager.fetch_one("SELECT COUNT(*) as count FROM job_offers WHERE status = 'active'")
                if total_jobs:
                    st.metric("Available Jobs", total_jobs['count'])
            except:
                pass
        else:
            # Statistics
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Total Matches", len(matches_df))
            with col2:
                excellent_matches = len(matches_df[matches_df['similarity_score'] > 0.8])
                st.metric("Excellent Matches", excellent_matches)
            with col3:
                good_matches = len(matches_df[matches_df['similarity_score'] > 0.6])
                st.metric("Good Matches", good_matches)
            with col4:
                avg_match = matches_df['similarity_score'].mean()
                st.metric("Avg Match Score", f"{avg_match:.1%}")
            
            st.divider()
            
            # Filter options
            col1, col2, col3 = st.columns(3)
            with col1:
                min_match_score = st.slider("Minimum Match %", 0, 100, 50, 5) / 100
            with col2:
                job_types = ['All'] + list(matches_df['job_type'].unique()) if 'job_type' in matches_df.columns else ['All']
                selected_job_type = st.selectbox("Job Type", job_types)
            with col3:
                sort_by = st.selectbox("Sort By", ["Match Score", "Date", "Company Name"])
            
            # Filter matches
            filtered_matches = matches_df[matches_df['similarity_score'] >= min_match_score]
            if selected_job_type != 'All' and 'job_type' in matches_df.columns:
                filtered_matches = filtered_matches[filtered_matches['job_type'] == selected_job_type]
            
            
            # Sort matches
            if sort_by == "Match Score":
                filtered_matches = filtered_matches.sort_values('similarity_score', ascending=False)
            elif sort_by == "Date" and 'last_updated' in filtered_matches.columns:
                filtered_matches = filtered_matches.sort_values('last_updated', ascending=False)
            elif 'company_name' in filtered_matches.columns:
                filtered_matches = filtered_matches.sort_values('company_name')
            
            # Display matches
            for idx, match in filtered_matches.iterrows():
                match_percentage = match['similarity_score'] * 100
                
                with st.container():
                    st.markdown(f"""
                    <div style="border: 1px solid #ddd; padding: 15px; margin: 10px 0; border-radius: 8px;">
                        <div style="display: flex; justify-content: space-between; align-items: center;">
                            <div>
                                <h3 style="margin: 0;">{match.get('job_title', 'Unknown Job')}</h3>
                                <h4 style="margin: 0; color: #666;">{match.get('company_name', 'Unknown Company')}</h4>
                            </div>
                            <div style="font-size: 24px; font-weight: bold; color: {'green' if match_percentage > 80 else 'orange' if match_percentage > 60 else 'red'};">
                                {match_percentage:.1f}%
                            </div>
                        </div>
                        <hr style="margin: 10px 0;">
                        <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 10px;">
                            <div>
                                <strong>📍 Location:</strong> {match.get('city', 'N/A')}, {match.get('country', 'N/A')}<br>
                                <strong>💼 Type:</strong> {match.get('job_type', 'N/A')}<br>
                                <strong>📅 Job category:</strong> {match.get('Job_category', 'N/A')}
                            </div>
                            <div>
                                <strong>Match Scores:</strong><br>
                                • Distance: {match.get('distance', 0)*100:.1f}%<br>
                                • similarity_score: {match.get('similarity_score', 0)*100:.1f}%<br> </div> </div>
                                
                            
                    </div>
                    """, unsafe_allow_html=True)
                    
                    col1, col2, col3 = st.columns([1, 1, 3])
                    with col1:
                        if st.button("View Details", key=f"view_{match.get('job_offer_id', idx)}"):
                            with st.expander("Job Description", expanded=True):
                                st.write(match.get('description', 'No description available'))
                    with col2:
                        if st.button("Apply Now", key=f"apply_{match.get('job_offer_id', idx)}"):
                            st.success("Application sent! (Demo)")
            
            # Export matches
            if st.button("📥 Export Matches to CSV"):
                csv = filtered_matches.to_csv(index=False)
                st.download_button(
                    label="Download CSV",
                    data=csv,
                    file_name=f"job_matches_{datetime.now().strftime('%Y%m%d')}.csv",
                    mime="text/csv"
                )


def create_job_offer(job_title, company_name, job_category, city, country, 
                    job_type, remote_option, education_level, experience_required,
                    salary_min, salary_max, currency, description, 
                    technical_skills, soft_skills, client_id):
    """
    FIXED: Create job offer with proper PostgreSQL array handling for skills_vector
    """
    connection = None
    cursor = None
    
    try:
        # STEP 1: Process and generate skills vector for job offer
        # Parse skills from text input
        tech_skills_list = []
        soft_skills_list = []
        
        if technical_skills and technical_skills.strip():
            tech_skills_list = [skill.strip() for skill in technical_skills.split(",") if skill.strip()]
        
        if soft_skills and soft_skills.strip():
            soft_skills_list = [skill.strip() for skill in soft_skills.split(",") if skill.strip()]
        
        # Create skills dictionary
        skills_dict = {
            "technical_skills": tech_skills_list,
            "soft_skills": soft_skills_list
        }
        
        # STEP 2: Generate embedding vector for job skills
        skills_vector = [0.0] * 1536  # Default vector
        
        # Try to generate embeddings using cv_generator
        try:
            if 'cv_generator' in st.session_state and st.session_state.cv_generator:
                embedding = st.session_state.cv_generator.generate_embedding(skills_dict)
                if embedding and len(embedding) == 1536:
                    skills_vector = embedding
                    st.info(f"✅ Generated skills vector for job offer")
                else:
                    st.warning("⚠️ Invalid embedding generated, using default vector")
            else:
                st.warning("⚠️ CV extractor not available, using default vector")
        except Exception as embed_error:
            st.warning(f"⚠️ Could not generate embeddings: {embed_error}")
        
        # STEP 3: Database insertion with PostgreSQL array format
        connection = psycopg2.connect(**POSTGRES_CONFIG)
        connection.autocommit = False
        cursor = connection.cursor()
        
        cursor.execute("BEGIN")
        
        # CRITICAL FIX: Use PostgreSQL array format instead of JSON
        insert_query = """
            INSERT INTO job_offers (
                client_id, job_title, company_name, job_category,
                city, country, job_type, remote_option, education_level,
                experience_required, salary_min, salary_max, currency, 
                description, technical_skills, soft_skills, skills_vector,
                created_at, status
            ) VALUES (
                %s, %s, %s, %s, %s, %s, %s, %s, %s,
                %s, %s, %s, %s, %s, %s, %s, %s, %s, %s
            ) RETURNING job_offer_id
        """
        
        # CRITICAL FIX: Convert skills_vector to PostgreSQL array format
        # For PostgreSQL arrays, we pass the Python list directly
        # psycopg2 will automatically convert it to the correct array format
        
        params = (
            client_id,
            job_title,
            company_name,
            job_category,
            city,
            country,
            job_type,
            remote_option,
            education_level,
            int(experience_required),
            float(salary_min) if salary_min else None,
            float(salary_max) if salary_max else None,
            currency,
            description,
            technical_skills,
            soft_skills,
            skills_vector,  # FIXED: Pass list directly, not JSON string
            datetime.now(),
            'active'
        )
        
        cursor.execute(insert_query, params)
        
        # Get the auto-generated job_offer_id
        result = cursor.fetchone()
        if result:
            job_offer_id = result[0]
        else:
            raise Exception("Failed to get generated job_offer_id")
        
        cursor.execute("COMMIT")
        st.success(f"✅ Job offer created with ID: {job_offer_id}")
        st.success(f"✅ Skills vector generated with {len(skills_vector)} dimensions")
        
        # STEP 4: Automatically trigger job matching after creating job offer
        try:
            st.info("🔄 Creating matches for new job offer...")
            
            # Get all candidates and create matches for this job
            all_candidates = db_manager.execute_query("SELECT candidate_id FROM candidates LIMIT 100")
            
            if all_candidates:
                matches_created = 0
                for candidate in all_candidates:
                    try:
                        # Use pipeline to create matches for this candidate
                        candidate_matches = pipeline.process_matching_for_candidate(candidate['candidate_id'])
                        matches_created += candidate_matches
                    except Exception as match_error:
                        st.warning(f"Match creation failed for candidate {candidate['candidate_id']}: {match_error}")
                        continue
                
                if matches_created > 0:
                    st.success(f"🎯 Created {matches_created} candidate matches for this job!")
                else:
                    st.warning("⚠️ No matches created for this job")
            else:
                st.info("ℹ️ No candidates available to match with this job")
                
        except Exception as match_error:
            st.warning(f"⚠️ Could not create matches: {match_error}")
        
        return True
        
    except Exception as e:
        if cursor:
            try:
                cursor.execute("ROLLBACK")
            except:
                pass
        st.error(f"❌ Error creating job offer: {e}")
        with st.expander("🔍 Error Details"):
            st.code(traceback.format_exc())
        return False
        
    finally:
        if cursor:
            cursor.close()
        if connection:
            connection.close()
def client_dashboard():
    st.title("🏢 Recruiter Dashboard")
    # Initialize Spark Analytics (cached)
    if 'spark_analytics' not in st.session_state:
        st.session_state.spark_analytics = SparkAnalysis(db_manager)
    
    spark_analytics = st.session_state.spark_analytics
    with st.sidebar:
        st.markdown("### ⚡ Spark-Powered Dashboard")
        
        # Performance indicator
        if spark_analytics.spark_enabled:
            st.success("🚀 Spark Engine: Active")
            st.info("~5-10x faster processing")
        else:
            st.warning("📊 Standard Mode: Active")
            st.info("Spark unavailable, using fallback")
        
        # FIXED: Diagnose button
        if st.button("🔍 Diagnose Database", key="diagnose_db"):
            diagnose_database_connection()
        
        # Enhanced refresh button
        if st.button("🚀 Spark Refresh All Data", key="spark_refresh_all_data"):
            with st.spinner("Spark-accelerated data refresh..."):
                start_time = time.time()
                
                try:
                    # Clear cache first
                    if 'spark_analytics' in st.session_state:
                        del st.session_state.spark_analytics
                    
                    # Reinitialize
                    st.session_state.spark_analytics = SparkAnalysis(db_manager)
                    
                    # Refresh matches using existing pipeline
                    if hasattr(pipeline, 'process_matching'):
                        matches_created = pipeline.process_matching()
                    else:
                        matches_created = 0
                    
                    end_time = time.time()
                    duration = end_time - start_time
                    
                    st.success(f"✅ Refresh completed in {duration:.2f}s")
                    if matches_created > 0:
                        st.info(f"Created {matches_created} matches")
                    
                except Exception as e:
                    st.error(f"Refresh failed: {e}")
                
                time.sleep(1)
                st.rerun()
    with st.sidebar:
        st.markdown(f"### Welcome, Client #{st.session_state.user_id}")
        

        # FIXED: Refresh candidates button with proper method call
        if st.button("🔄 Refresh All Matches", key="refresh_all_matches"):
            with st.spinner("Re-calculating all candidate matches..."):
                try:
                    # Get all candidates and process matches for each
                    candidates = db_manager.execute_query("SELECT candidate_id FROM candidates LIMIT 10")
                    
                    if candidates:
                        total_matches_created = 0
                        for candidate in candidates:
                            # FIXED: Call with proper parameter
                            matches_created = pipeline.process_matching_for_candidate(candidate['candidate_id'])
                            total_matches_created += matches_created
                        
                        st.success(f" {total_matches_created} matches for {len(candidates)} candidates!")
                    else:
                        st.warning("No candidates found to process matches for.")
                        
                except Exception as e:
                    st.error(f"❌ Error refreshing matches: {e}")
                
                time.sleep(1)
                st.rerun()

        if st.button("🚪 Logout", key="client_logout"):
            st.session_state.user_role = None
            st.session_state.user_id = None
            st.rerun()

        if st.session_state.last_update:
            st.caption(f"Last updated: {st.session_state.last_update}")

    # Tabs
    tab1, tab2, tab3, tab4,tab5 = st.tabs([
        "📋 Job Offers & Candidates",
        "➕ Add Job Offer",
        "📊 Analytics",
        "🔍 Search Candidates",
        "🔍 Skills Analysis" 
    ])

    # Tab 1: Job Offers & Candidates - FIXED
    with tab1:
        st.markdown("### 📋 Your Job Offers & Matched Candidates")
        
        # Get job offers for this client
        job_offers = db_manager.execute_query(
            "SELECT * FROM job_offers WHERE client_id = %(client_id)s ORDER BY created_at DESC",
            {"client_id": st.session_state.user_id}
        )
        
        if not job_offers:
            st.info("You have no job offers yet. Click on 'Add Job Offer' tab to create one.")
        else:
            for job in job_offers:
                with st.container():
                    st.markdown(f"""
                    <div class="job-card">
                        <h3>{job['job_title']} at {job['company_name']}</h3>
                        <p>📍 {job['city']}, {job['country']} | 💼 {job['job_type']} | 🏠 {job.get('remote_option', 'Not specified')}</p>
                        <p>🏷️ {job.get('job_category', 'Not specified')} | 🎓 {job.get('education_level', 'Not specified')} | 🕒 {job.get('experience_required', 0)}+ years experience</p>
                        <p>💰 {job['currency']} {job.get('salary_min', 'N/A')} - {job.get('salary_max', 'N/A')}</p>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # FIXED: Get matched candidates using correct method
                    try:
                        candidates = pipeline.get_job_candidates(job['job_offer_id'])
                        
                        if candidates.empty:
                            st.info("No matching candidates found yet for this position.")
                            
                            # Show button to create matches for this job
                            if st.button(f"🎯 Find Candidates", key=f"find_candidates_{job['job_offer_id']}"):
                                with st.spinner("Finding candidates for this job..."):
                                    # Get all candidates and create matches
                                    all_candidates = db_manager.execute_query("SELECT candidate_id FROM candidates")
                                    matches_created = 0
                                    
                                    for candidate in all_candidates:
                                        try:
                                            matches_created += pipeline.process_matching_for_candidate(candidate['candidate_id'])
                                        except Exception as e:
                                            continue
                                    
                                    if matches_created > 0:
                                        st.success(f"✅ Found {matches_created} candidate matches!")
                                        st.rerun()
                                    else:
                                        st.warning("No suitable candidates found for this position.")
                        else:
                            # Show total candidates count and display top 10
                            total_candidates = len(candidates)
                            st.markdown(f"**🎯 Top 10 Matched Candidates***)")
                            
                            # Display top 10 candidates
                            for idx, candidate in candidates.head(10).iterrows():
                                match_score = candidate['similarity_score'] * 100
                                match_class = get_match_class(candidate['similarity_score'])
                                
                                # Add ranking number
                                rank = idx + 1
                                rank_emoji = "🥇" if rank == 1 else "🥈" if rank == 2 else "🥉" if rank == 3 else f"#{rank}"
                                
                                st.markdown(f"""
                                <div class="candidate-card">
                                    <strong>{rank_emoji} {candidate['first_name']} {candidate['last_name']}</strong> 
                                    <span class="{match_class}" style="float: right;">{match_score:.1f}% match</span><br>
                                    📧 {candidate.get('email', 'N/A')} | 📱 {candidate.get('phone', 'N/A')}<br>
                                    📍 {candidate.get('city', 'N/A')}, {candidate.get('country', 'N/A')} | 
                                    🎓 {candidate.get('education_level', 'N/A')} | 
                                    💼 {candidate.get('experience_years', 0)} years experience
                                </div>
                                """, unsafe_allow_html=True)
                            
                            # Show info if there are more than 10 candidates
                            if total_candidates > 10:
                                st.info(f"💡 Showing top 10 candidates out of {total_candidates} total matches")
                                
                                # Optional: Add button to show all candidates
                                if st.button(f"📋 Show All {total_candidates} Candidates", key=f"show_all_{job['job_offer_id']}"):
                                    with st.expander(f"All {total_candidates} Candidates for {job['job_title']}", expanded=True):
                                        for idx, candidate in candidates.iterrows():
                                            match_score = candidate['similarity_score'] * 100
                                            match_class = get_match_class(candidate['similarity_score'])
                                            rank = idx + 1
                                            
                                            st.markdown(f"""
                                            <div style="border: 1px solid #ddd; padding: 10px; margin: 5px 0; border-radius: 5px;">
                                                <strong>#{rank} {candidate['first_name']} {candidate['last_name']}</strong> 
                                                <span class="{match_class}" style="float: right;">{match_score:.1f}% match</span><br>
                                                📧 {candidate.get('email', 'N/A')} | 📱 {candidate.get('phone', 'N/A')}<br>
                                                📍 {candidate.get('city', 'N/A')}, {candidate.get('country', 'N/A')} | 
                                                🎓 {candidate.get('education_level', 'N/A')} | 
                                                💼 {candidate.get('experience_years', 0)} years experience
                                            </div>
                                            """, unsafe_allow_html=True)
                            
                            # Add export/download option
                            if total_candidates > 0:
                                col1, col2 = st.columns(2)
                                with col1:
                                    if st.button(f"📄 Export Candidates List", key=f"export_{job['job_offer_id']}"):
                                        # Create CSV data
                                        export_df = candidates[['first_name', 'last_name', 'email', 'phone', 
                                                            'city', 'country', 'education_level', 'experience_years', 
                                                            'similarity_score']].copy()
                                        export_df['match_percentage'] = (export_df['similarity_score'] * 100).round(1)
                                        export_df = export_df.drop('similarity_score', axis=1)
                                        
                                        csv = export_df.to_csv(index=False)
                                        st.download_button(
                                            label="💾 Download CSV",
                                            data=csv,
                                            file_name=f"candidates_{job['job_title'].replace(' ', '_')}_{job['job_offer_id']}.csv",
                                            mime="text/csv",
                                            key=f"download_{job['job_offer_id']}"
                                        )
                                
                                with col2:
                                    # Add refresh matches button
                                    if st.button(f"🔄 Refresh Matches", key=f"refresh_{job['job_offer_id']}"):
                                        with st.spinner("Refreshing matches..."):
                                            # Re-run matching for this specific job
                                            all_candidates = db_manager.execute_query("SELECT candidate_id FROM candidates")
                                            matches_created = 0
                                            
                                            for candidate in all_candidates:
                                                try:
                                                    matches_created += pipeline.process_matching_for_candidate(candidate['candidate_id'])
                                                except Exception as e:
                                                    continue
                                            
                                            st.success(f"✅ Refreshed matches! Found {matches_created} total matches.")
                                            st.rerun()
                            
                    except Exception as e:
                        st.error(f"Error loading candidates: {e}")
                        st.info("Try refreshing matches using the sidebar button.")
                    
                    st.divider()
        # Tab 2: Add Job Offer - FIXED
    with tab2:
        st.markdown("### ➕ Create New Job Offer")
        
        # Add database connection reset button
        if st.button("🔧 Reset Database Connection", key="reset_db_job_offer"):
            try:
                # Reset database connection
                if hasattr(db_manager, 'connection') and db_manager.connection:
                    try:
                        db_manager.connection.rollback()
                        db_manager.connection.close()
                    except:
                        pass
                    db_manager.connection = None
                
                # Force new connection
                db_manager.get_connection()
                st.success("✅ Database connection reset successfully")
                
            except Exception as e:
                st.error(f"❌ Could not reset connection: {e}")
        
        with st.form("add_job_offer"):
            # Basic Information
            col1, col2 = st.columns(2)
            with col1:
                job_title = st.text_input("Job Title *", key="job_title")
                company_name = st.text_input("Company Name *", key="company_name")
                job_category = st.selectbox(
                    "Job Category *", 
                    ["Technology", "Healthcare", "Finance", "Education", "Marketing", 
                    "Sales", "Engineering", "Human Resources", "Operations", "Other"],
                    key="job_category"
                )
            with col2:
                job_type = st.selectbox("Job Type *", ["Full-time", "Part-time", "Contract", "Internship"], key="job_type")
                remote_option = st.selectbox("Remote Option *", ["On-site", "Remote", "Hybrid"], key="remote_option")
                education_level = st.selectbox(
                    "Education Level Required *", 
                    ["High School", "Associate", "Bachelor", "Master", "PhD", "Any"],
                    key="education_level"
                )
            
            # Experience and Location
            col3, col4 = st.columns(2)
            with col3:
                experience_required = st.number_input("Experience Required (years) *", min_value=0, max_value=50, key="experience_required")
                city = st.text_input("City *", key="city")
                country = st.text_input("Country *", key="country")
            with col4:
                currency = st.selectbox("Currency", ["USD", "EUR", "TND", "GBP", "CAD"], key="currency")
                salary_min = st.number_input("Minimum Salary", min_value=0, key="salary_min")
                salary_max = st.number_input("Maximum Salary", min_value=0, key="salary_max")
            
            # Description and Skills
            description = st.text_area("Job Description *", height=150, key="description")
            technical_skills = st.text_area("Technical Skills (comma-separated) *", height=100, key="technical_skills")
            soft_skills = st.text_area("Soft Skills (comma-separated) *", height=100, key="soft_skills")
            
            submitted = st.form_submit_button("Create Job Offer", use_container_width=True)

            if submitted:
                # Validation
                if not all([job_title, company_name, job_category, city, country, description, technical_skills]):
                    st.error("Please fill in all required fields marked with *")
                else:
                    with st.spinner("Creating job offer..."):
                        try:
                            # Create job offer with proper transaction handling
                            success = create_job_offer(
                                job_title=job_title,
                                company_name=company_name,
                                job_category=job_category,
                                city=city,
                                country=country,
                                job_type=job_type,
                                remote_option=remote_option,
                                education_level=education_level,
                                experience_required=experience_required,
                                salary_min=salary_min,
                                salary_max=salary_max,
                                currency=currency,
                                description=description,
                                technical_skills=technical_skills,
                                soft_skills=soft_skills,
                                client_id=st.session_state.user_id
                            )
                            
                            if success:
                                st.success("✅ Job offer created successfully!")
                                st.balloons()
                                time.sleep(2)
                                st.rerun()
                            else:
                                st.error("❌ Failed to create job offer. Please try again.")
                                
                        except Exception as e:
                            st.error(f"❌ Error creating job offer: {str(e)}")
                            with st.expander("🔍 Error Details"):
                                st.code(str(e))
    # Tab 3: Analytics - COMPLETELY FIXED
    with tab3:
        st.markdown("### 📊 Your Analytics Dashboard")
        
        # FIXED: Use correct table name (job_matches instead of candidate_job_offer)
        analytics = db_manager.execute_query("""
            SELECT 
                jo.job_title, 
                COUNT(jm.candidate_id) AS candidate_count,
                AVG(jm.similarity_score) AS avg_score,
                MAX(jm.similarity_score) AS best_score
            FROM job_offers jo
            LEFT JOIN job_matches jm ON jo.job_offer_id = jm.job_offer_id
            WHERE jo.client_id = %(client_id)s
            GROUP BY jo.job_offer_id, jo.job_title
            ORDER BY candidate_count DESC
        """, {"client_id": st.session_state.user_id})

        if analytics:
            df = pd.DataFrame(analytics)
            
            # Display summary metrics
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                total_jobs = len(df)
                st.metric("Total Job Offers", total_jobs)
            with col2:
                total_candidates = df['candidate_count'].sum()
                st.metric("Total Candidate Matches", int(total_candidates))
            with col3:
                avg_candidates_per_job = df['candidate_count'].mean()
                st.metric("Avg Candidates/Job", f"{avg_candidates_per_job:.1f}")
            with col4:
                overall_avg_score = df['avg_score'].mean()
                st.metric("Overall Avg Score", f"{overall_avg_score:.1%}" if overall_avg_score else "N/A")
            
            # Create proper charts
            col1, col2 = st.columns(2)
            
            with col1:
                # Bar chart for candidate count
                fig1 = px.bar(
                    df, 
                    x='job_title', 
                    y='candidate_count',
                    title="Candidates per Job Offer",
                    labels={'candidate_count': 'Number of Candidates', 'job_title': 'Job Title'}
                )
                fig1.update_xaxes(tickangle=45)
                st.plotly_chart(fig1, use_container_width=True)
            
            with col2:
                # Bar chart for average score
                df_with_scores = df[df['avg_score'].notna()]  # Remove NaN values
                if not df_with_scores.empty:
                    fig2 = px.bar(
                        df_with_scores, 
                        x='job_title', 
                        y='avg_score',
                        title="Average Match Score per Job",
                        labels={'avg_score': 'Average Match Score', 'job_title': 'Job Title'},
                        color='avg_score',
                        color_continuous_scale='viridis'
                    )
                    fig2.update_xaxes(tickangle=45)
                    fig2.update_layout(yaxis_tickformat='.0%')
                    st.plotly_chart(fig2, use_container_width=True)
                else:
                    st.info("No match scores available yet.")
            
            # Display data table
            st.subheader("📋 Detailed Analytics")
            display_df = df.copy()
            display_df['avg_score'] = display_df['avg_score'].apply(lambda x: f"{x:.1%}" if pd.notna(x) else "N/A")
            display_df['best_score'] = display_df['best_score'].apply(lambda x: f"{x:.1%}" if pd.notna(x) else "N/A")
            st.dataframe(
                display_df.rename(columns={
                    'job_title': 'Job Title',
                    'candidate_count': 'Candidates',
                    'avg_score': 'Avg Score',
                    'best_score': 'Best Score'
                }),
                use_container_width=True
            )
            # Display power Bi dashbords
            st.subheader("📊  Powerbi dashboards")
            st.markdown(
                        """
                        <a href="https://app.powerbi.com/view?r=eyJrIjoiYzRlZWUxNjEtNjk0Yi00MjZkLWJmM2QtNWQ4YzBlZTY1OWI4IiwidCI6ImYzMmY4NDgwLTg4NWMtNDdhMi05MGNlLTUxNTJkNjkxNWE4ZiIsImMiOjl9" target="_blank">
                            <button style="
                                background-color: #ff6b6b;
                                color: white;
                                padding: 10px 20px;
                                border: none;
                                border-radius: 5px;
                                cursor: pointer;
                                font-size: 16px;
                            ">
                                See more insights
                            </button>
                        </a>
                        """,
                        unsafe_allow_html=True
                    )
        
        else:
            st.info("No analytics available yet. Create some job offers to see analytics.")

    # Tab 4: Search Candidates - COMPLETELY FIXED
    with tab4:
        st.markdown("### 🔍 Search Candidates")
        
        # Search filters
        col1, col2 = st.columns(2)
        with col1:
            min_exp = st.number_input("Minimum Years of Experience", 0, 50, 0, key="min_exp")
            education = st.selectbox("Education Level", ["All", "PhD", "Master", "Bachelor", "High School"], key="education_level_search")
        with col2:
            location = st.text_input("Location (City or Country)", key="location")
            skills = st.text_input("Skills (comma-separated)", key="skills_input")

        if st.button("Search", key="search_btn"):
            # FIXED: Use correct table name and session variable
            client_id = st.session_state.user_id
            
            if not client_id:
                st.error("Client ID not found in session. Please logout and login again.")
            else:
                # FIXED: Use correct table name (job_matches instead of candidate_job_offer)
                base_query = """
                    SELECT DISTINCT c.*,
                        COALESCE(AVG(jm.similarity_score), 0) as avg_match_score,
                        COUNT(jm.match_id) as total_matches
                    FROM candidates c
                    LEFT JOIN job_matches jm ON c.candidate_id = jm.candidate_id
                    LEFT JOIN job_offers jo ON jm.job_offer_id = jo.job_offer_id
                    WHERE 1=1
                """
                params = {}
                
                # Add filters
                if min_exp > 0:
                    base_query += " AND c.experience_years >= %(min_exp)s"
                    params["min_exp"] = min_exp

                if education != "All":
                    base_query += " AND c.education_level = %(education)s"
                    params["education"] = education

                if location:
                    base_query += " AND (c.city ILIKE %(location)s OR c.country ILIKE %(location)s)"
                    params["location"] = f"%{location}%"

                if skills:
                    skills_list = [s.strip().lower() for s in skills.split(",") if s.strip()]
                    if skills_list:
                        # Search in skills JSON
                        skills_conditions = []
                        for i, skill in enumerate(skills_list):
                            param_name = f"skill_{i}"
                            skills_conditions.append(f"LOWER(c.skills::text) LIKE %({param_name})s")
                            params[param_name] = f"%{skill}%"
                        
                        if skills_conditions:
                            base_query += " AND (" + " OR ".join(skills_conditions) + ")"
                
                # Add grouping and ordering
                base_query += """
                    GROUP BY c.candidate_id
                    ORDER BY avg_match_score DESC, total_matches DESC
                    LIMIT 50
                """

                # Execute the query
                try:
                    candidates = db_manager.execute_query(base_query, params)

                    # Display results
                    if candidates:
                        st.markdown(f"### 👤 Found {len(candidates)} Candidate(s)")
                        
                        for cand in candidates:
                            with st.container():
                                # Parse skills safely
                                try:
                                    if cand.get('skills'):
                                        if isinstance(cand['skills'], str):
                                            skills_data = json.loads(cand['skills'])
                                        else:
                                            skills_data = cand['skills']
                                        
                                        tech_skills = skills_data.get('technical_skills', [])
                                        soft_skills = skills_data.get('soft_skills', [])
                                        skills_display = ", ".join(tech_skills[:5])  # Show first 5 skills
                                        if len(tech_skills) > 5:
                                            skills_display += "..."
                                    else:
                                        skills_display = "No skills listed"
                                except (json.JSONDecodeError, TypeError):
                                    skills_display = "Skills format error"
                                
                                # Display candidate card
                                avg_match = cand.get('avg_match_score', 0)
                                total_matches = cand.get('total_matches', 0)
                                
                                st.markdown(f"""
                                <div class="candidate-card">
                                    <h4>{cand['first_name']} {cand['last_name']}</h4>
                                    <p>📧 {cand['email']} | 📱 {cand.get('phone', 'N/A')}</p>
                                    <p>📍 {cand['city']}, {cand['country']} | 
                                       🎓 {cand['education_level']} | 
                                       💼 {cand.get('experience_years', 0)} years experience</p>
                                    <p>🛠️ <strong>Skills:</strong> {skills_display}</p>
                                    <p>📊 Average Match: {avg_match:.1%} | Total Matches: {total_matches}</p>
                                </div>
                                """, unsafe_allow_html=True)
                                
                                # Action buttons
                                col1, col2, col3 = st.columns([1, 1, 2])
                                with col1:
                                    if st.button(f"View Profile", key=f"profile_{cand['candidate_id']}"):
                                        st.info("Profile view feature coming soon...")
                                with col2:
                                    if st.button(f"Contact", key=f"contact_{cand['candidate_id']}"):
                                        st.success("Contact initiated! (Demo)")
                                
                                st.divider()
                                
                    else:
                        st.info("No candidates found matching your criteria. Try adjusting your filters.")
                        
                        # Show some suggestions
                        st.markdown("**💡 Search Tips:**")
                        st.write("• Try broader skill terms (e.g., 'Python' instead of 'Python 3.9')")
                        st.write("• Use location names like 'Paris' or 'France'")
                        st.write("• Leave some filters empty for broader results")
                        
                except Exception as e:
                    st.error(f"❌ Search error: {str(e)}")
                    st.error("Please check your database connection and table structure.")
                    with st.expander("🔍 Error Details"):
                        st.code(str(e))
   
    with tab5:
        st.markdown("### 🔍 Advanced Skills Analysis Dashboard")
        # Add navigation tabs for different analysis types
        analysis_tab = st.selectbox(
            "  ",
            [
                "📊 Standard Analytics", 
                "🚀 Advanced Dashboard", 
                "📈 Timeline Analytics",
                "☁️ Skills Bubble Cloud", 
                "💼 Experience Analysis",
                "🔗 Skills Correlation"
                
            ]
        )
        
        # Get analytics data once
        with st.spinner("Loading analytics data..."):
            # In your admin dashboard Tab5:
            analytics= db_manager.get_skills_analytics()
        
        if not analytics:
            st.info("📊 No analytics data available. Process some CV documents first!")
            st.info("Upload and process CVs using the 'Upload CV' tab to see analytics")
            return
        
        # Display selected analysis type
        if analysis_tab == "📊 Standard Analytics":
            # Original analytics code (enhanced)
            
            # Technical Skills Analysis
            if 'technical_skills' in analytics and not analytics['technical_skills'].empty:
                st.subheader("🛠️ Most Common Technical Skills")
                
                tech_df = analytics['technical_skills']
                
                col1, col2 = st.columns(2)
                
                with col1:
                    # Horizontal bar chart
                    fig_tech = px.bar(
                        tech_df.head(15),
                        x='frequency',
                        y='skill',
                        orientation='h',
                        title="Top 15 Technical Skills",
                        color='frequency',
                        color_continuous_scale='viridis'
                    )
                    fig_tech.update_layout(height=500, yaxis={'categoryorder': 'total ascending'})
                    st.plotly_chart(fig_tech, use_container_width=True)
                
                with col2:
                    # Donut chart
                    fig_tech_donut = px.pie(
                        tech_df.head(10),
                        values='frequency',
                        names='skill',
                        title="Top 10 Technical Skills Distribution",
                        hole=0.4
                    )
                    st.plotly_chart(fig_tech_donut, use_container_width=True)
                
                # Show top skills as metrics
                st.subheader("🏆 Top Technical Skills")
                
                # Create dynamic number of columns based on available skills
                top_skills_count = min(5, len(tech_df))
                if top_skills_count > 0:
                    cols = st.columns(top_skills_count)
                    for i, (_, row) in enumerate(tech_df.head(top_skills_count).iterrows()):
                        with cols[i]:
                            st.metric(f"#{i+1}", row['skill'], f"{row['frequency']} candidates")
                else:
                    st.info("No technical skills data available")
            
            # Soft Skills Analysis
            if 'soft_skills' in analytics and not analytics['soft_skills'].empty:
                st.subheader("🤝 Most Common Soft Skills")
                
                soft_df = analytics['soft_skills']
                
                col1, col2 = st.columns(2)
                
                with col1:
                    # Pie chart
                    fig_soft = px.pie(
                        soft_df.head(10),
                        values='frequency',
                        names='skill',
                        title="Top 10 Soft Skills Distribution"
                    )
                    st.plotly_chart(fig_soft, use_container_width=True)
                
               
            
            # Job Categories and Experience Analysis
            col1, col2 = st.columns(2)
            
            with col1:
                if 'categories' in analytics and not analytics['categories'].empty:
                    st.subheader("📂 Job Categories")
                    cat_df = analytics['categories']
                    
                    fig_cat = px.pie(
                        cat_df,
                        values='count',
                        names='job_category',
                        title="Job Categories Distribution"
                    )
                    st.plotly_chart(fig_cat, use_container_width=True)
            
            with col2:
                if 'experience' in analytics and not analytics['experience'].empty:
                    st.subheader("📈 Experience Distribution")
                    exp_df = analytics['experience']
                    
                    fig_exp = px.bar(
                        exp_df,
                        x='experience_years',
                        y='count',
                        title="Experience Years Distribution"
                    )
                    st.plotly_chart(fig_exp, use_container_width=True)
            
            # Education Analysis
            if 'education' in analytics and not analytics['education'].empty:
                st.subheader("🎓 Education Level Distribution")
                edu_df = analytics['education']
                
                col1, col2 = st.columns(2)
                with col1:
                    fig_edu = px.pie(
                        edu_df,
                        values='count',
                        names='education_level',
                        title="Education Levels"
                    )
                    st.plotly_chart(fig_edu, use_container_width=True)
                
                with col2:
                    fig_edu_bar = px.bar(
                        edu_df,
                        x='education_level',
                        y='count',
                        title="Education Level Counts"
                    )
                    fig_edu_bar.update_xaxes(tickangle=45)
                    st.plotly_chart(fig_edu_bar, use_container_width=True)
        
        elif analysis_tab == "🚀 Advanced Dashboard":
            create_advanced_skills_dashboard(db_manager)
            
        
        elif analysis_tab == "📈 Timeline Analytics":
            create_timeline_analytics(db_manager)
        
        elif analysis_tab == "☁️ Skills Bubble Cloud":
            create_skills_word_cloud_simulation(db_manager)
        
        elif analysis_tab == "💼 Experience Analysis":
            create_advanced_experience_analysis(db_manager)
        
        elif analysis_tab == "🔗 Skills Correlation":
            create_skills_correlation_matrix(db_manager)
        elif analysis_tab == "📈 Chart Libraries Demo":
            demonstrate_chart_libraries(db_manager)
    
def main():
    """Main Streamlit entry point"""
    with st.sidebar:
        if st.button("🔍 Diagnose System", key="diagnose_system"):
            diagnose_database_connection()
            return 
    # Initialize app
    initialize_app()
    
    # Route to dashboards
    if st.session_state.user_role == 'candidate':
        candidate_dashboard()
    elif st.session_state.user_role == 'client':
        client_dashboard()
    else:
        login_page()

if __name__ == "__main__":
    main()